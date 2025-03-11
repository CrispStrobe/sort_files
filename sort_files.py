#!/usr/bin/env python3
import os
import sys
import re
import argparse
import subprocess
import shutil
import tempfile
import logging
import logging.handlers
import queue
from datetime import datetime
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
import signal
import platform
import time
from types import MappingProxyType

# Import tqdm for progress bar
try:
    import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    print("WARNING: tqdm is not installed. Install it using 'pip install tqdm' to see progress bars.")

# Global lock for thread-safe file operations
file_lock = threading.Lock()
thread_local = threading.local()

# Configuration Constants
MODEL_NAME = "cas/spaetzle-v85-7b"  
# "cas/llama-3.2-3b-instruct"
# "cas/spaetzle-v85-7b"  
# Change to other ollama model as fits your needs. Ensure this is defined before usage

# Timeout settings
OCR_PROCESS_TIMEOUT = 120  # 2 minutes timeout for OCR processes
DEFAULT_PROCESS_TIMEOUT = 60  # 1 minute timeout for other processes
MAX_WAIT_AFTER_SHUTDOWN = 10  # Maximum seconds to wait after shutdown signal

MAX_RETRY_ATTEMPTS = 1
SUPPORTED_EXTENSIONS = (
    '.csv', '.doc', '.docx', '.eml', '.epub', '.gif', '.jpg', '.jpeg', '.json', 
    '.html', '.htm', '.mp3', '.msg', '.odt', '.ogg', '.pdf', '.png', '.pptx', 
    '.ps', '.rtf', '.tiff', '.tif', '.txt', '.wav', '.xlsx', '.xls', '.ppt', 
    '.ods', '.odp', '.md'
)
OCR_PREFERRED_ORDER = ['tesseract', 'kraken']  # Define preferred OCR methods

# Define required Tesseract languages
REQUIRED_TESS_LANGS = ['eng', 'deu', 'ara']

# Define a mapping from languages/scripts to Kraken model IDs
LANGUAGE_MODEL_MAPPING = {
    'en': 'en_best.mlmodel',
    'de': 'german_print.mlmodel',
    'ar': 'arabic_best.mlmodel',
    'de_np': 'austriannewspapers.mlmodel',
    'lat': 'reichenau_lat_cat_099218.mlmodel',
    'lat_cr': ' cremma-generic-1.0.1.mlmodel',
    # Add other mappings as needed
}

class TqdmLoggingHandler(logging.Handler):
    """
    Custom logging handler that uses tqdm.write() to output log messages.
    This ensures that log messages do not interfere with the tqdm progress bar.
    """
    def emit(self, record):
        try:
            msg = self.format(record)
            if hasattr(tqdm, 'write'):
                tqdm.tqdm.write(msg)
            else:
                print(msg)  # Fallback to print if tqdm.write is not available
        except Exception:
            self.handleError(record)


class ProgressBar:
    """
    A class to handle progress bar operations using tqdm.
    
    If tqdm is unavailable, it gracefully degrades without displaying a progress bar.
    """
    def __init__(self, total, desc="Processing Files", enable=True):
        self.enable = enable
        self.lock = threading.Lock()
        self.bar = None
        if self.enable and TQDM_AVAILABLE:
            try:
                self.bar = tqdm.tqdm(total=total, desc=desc, unit="file")
                logging.debug("Progress bar initialized using tqdm.")
            except Exception as e:
                logging.error(f"Failed to initialize progress bar: {e}")
                self.enable = False
        else:
            logging.warning("tqdm is not available or progress bar is disabled.")

    def update(self, n=1):
        if self.enable and self.bar:
            with self.lock:
                self.bar.update(n)
    
    def close(self):
        if self.enable and self.bar:
            with self.lock:
                self.bar.close()
                logging.debug("Progress bar closed.")

def is_file_empty(file_path):
    return os.path.getsize(file_path) == 0

def get_kraken_model(language='en'):
    """
    Retrieve the Kraken model based on the language code.
    
    Parameters:
        language (str): Language code (e.g., 'en', 'de', 'ar').
    
    Returns:
        str: The Kraken model filename to use.
    """
    # Get the model filename based on language (default to 'en_best.mlmodel')
    model_filename = LANGUAGE_MODEL_MAPPING.get(language, 'en_best.mlmodel')
    
    # Get the Kraken model path from environment variable (if set)
    model_filepath = os.getenv('KRAKEN_MODEL_PATH', '')

    # If a model path is provided, build the full path, otherwise return just the model name
    if model_filepath:
        model_path = os.path.expanduser(f'{model_filepath}/{model_filename}')
        # Check if the file exists at the specified location
        if not os.path.isfile(model_path):
            logging.error(f"Kraken model '{model_filename}' not found at {model_filepath}.")
        return model_path
    else:
        # No path provided, return just the model filename and rely on Kraken's internal search
        logging.info(f"Using Kraken model '{model_filename}' from default location.")
        return model_filename

def get_openai_client():
    if not hasattr(thread_local, "client"):
        try:
            thread_local.client = OpenAI(
                base_url="http://localhost:11434/v1/",
                api_key="ollama"
            )
            logging.debug("OpenAI client initialized successfully for thread.")
        except Exception as e:
            logging.critical(f"Failed to initialize OpenAI client in thread: {e}")
            raise
    return thread_local.client

def setup_logging(verbose, debug):
    log_format = '%(asctime)s - %(levelname)s - %(message)s'
    
    # Create a log queue
    log_queue = queue.Queue(-1)  # No limit on size

    # Create a formatter
    formatter = logging.Formatter(log_format)

    # Create log directory if it doesn't exist
    if not os.path.exists('logs'):
        os.makedirs('logs')

    # Determine log file and level based on flags
    # NOTSET=0 DEBUG=10 INFO=20 WARN=30 ERROR=40 CRITICAL=50
    if debug:
        log_filename = os.path.join('logs', 'script_debug.log')
        file_level = logging.DEBUG
        stream_level = logging.DEBUG
    elif verbose:
        log_filename = os.path.join('logs', 'script_verbose.log')
        file_level = logging.INFO
        stream_level = logging.INFO
    else:
        log_filename = os.path.join('logs', 'script.log')
        file_level = logging.INFO
        stream_level = logging.CRITICAL  # **Only CRITICAL logs are shown to console by default**

    # File handler
    file_handler = logging.FileHandler(log_filename, mode='a')
    file_handler.setLevel(file_level)
    file_handler.setFormatter(formatter)

    # Stream handler using TqdmLoggingHandler
    stream_handler = TqdmLoggingHandler()
    stream_handler.setLevel(stream_level)
    stream_handler.setFormatter(formatter)

    # Create logger
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG if debug else (logging.INFO if verbose else logging.ERROR))

    # Remove existing handlers
    if logger.hasHandlers():
        logger.handlers.clear()

    # Add handlers
    logger.addHandler(file_handler)
    logger.addHandler(stream_handler)

    # Suppress external library logs unless in debug mode
    external_loggers = [
        'openai', 'pdfminer', 'PIL', 'urllib3', 'requests', 
        'tqdm', 'textract'
    ]
    for lib in external_loggers:
        lib_logger = logging.getLogger(lib)
        if debug:
            lib_logger.setLevel(logging.DEBUG)
        elif verbose:
            lib_logger.setLevel(logging.CRITICAL)
        else:
            lib_logger.setLevel(logging.CRITICAL)  # **Suppress ERROR logs by setting to CRITICAL**
    
    logging.debug("Logging is set up successfully.")


def prompt_install_tool(tool_name, install_instructions, auto_install=False, no_install=False):
    """
    Prompt the user to install a missing tool.

    Parameters:
        tool_name (str): The name of the tool to install.
        install_instructions (list): The command to install the tool.
        auto_install (bool): If True, attempt to install without prompting.
        no_install (bool): If True, do not attempt to install even if prompted.

    Returns:
        bool: True if installation succeeded, False otherwise.
    """
    if no_install:
        logging.info(f"Installation of {tool_name} is disabled via --no-install flag.")
        return False
    if auto_install:
        user_input = 'y'
    else:
        try:
            user_input = input(f"{tool_name} is not installed. Do you want to install it now? [y/N]: ").strip().lower()
        except EOFError:
            logging.warning(f"Input not available. Skipping installation of {tool_name}.")
            return False

    if user_input == 'y':
        try:
            subprocess.run(install_instructions, check=True)
            logging.info(f"Successfully installed {tool_name}.")
            return True
        except subprocess.CalledProcessError as e:
            logging.error(f"Failed to install {tool_name}: {e}")
    else:
        logging.info(f"{tool_name} is required for OCR functionalities. Proceeding without it.")
    return False


def find_tessdata_dirs():
    """
    Search for tessdata directories in common locations.
    Returns a list of found tessdata directories.
    """
    common_paths = [
        '/usr/local/share/tessdata',
        '/usr/share/tesseract-ocr/4.00/tessdata',
        '/usr/share/tesseract-ocr/tessdata',
        '/opt/homebrew/share/tessdata',
        '/usr/local/Cellar/tesseract/5.4.1/share/tessdata',  # Adjust based on your Tesseract version
    ]
    tessdata_dirs = []

    # First, check if TESSDATA_PREFIX is set in the environment
    env_prefix = os.environ.get('TESSDATA_PREFIX')
    if env_prefix and os.path.isdir(env_prefix):
        tessdata_dirs.append(env_prefix)
        logging.debug(f"Found tessdata directory from TESSDATA_PREFIX: {env_prefix}")
    
    # Search common paths if not found in environment
    if not tessdata_dirs:
        for path in common_paths:
            if os.path.isdir(path):
                tessdata_dirs.append(path)
                logging.debug(f"Found tessdata directory: {path}")
    
    # Optionally, search the entire system (commented out for performance)
    # for root, dirs, files in os.walk('/'):
    #     if 'tessdata' in dirs:
    #         tessdata_dir = os.path.join(root, 'tessdata')
    #         tessdata_dirs.append(tessdata_dir)
    #         logging.debug(f"Found tessdata directory via system-wide search: {tessdata_dir}")
    
    return tessdata_dirs

def check_tesseract_langs(required_langs):
    """
    Check if the required Tesseract language data files are present.
    """
    try:
        langs = subprocess.check_output(['tesseract', '--list-langs'], stderr=subprocess.STDOUT).decode('utf-8')
        missing_langs = [lang for lang in required_langs if lang not in langs]
        if missing_langs:
            logging.error(f"Missing Tesseract language data files: {', '.join(missing_langs)}")
            return False
        logging.debug("All required Tesseract language data files are present.")
        return True
    except subprocess.CalledProcessError as e:
        logging.error(f"Error checking Tesseract languages: {e}")
        return False
    except FileNotFoundError:
        logging.error("Tesseract is not installed or not found in PATH.")
        return False
    except Exception as e:
        logging.error(f"Unexpected error checking Tesseract languages: {e}")
        return False
    
def process_with_timeout(func, *args, **kwargs):
    """
    Run a function with a timeout.
    """
    result = None
    future = None
    
    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(func, *args, **kwargs)
            result = future.result(timeout=300)  # 5-minute timeout for each file processing
    except concurrent.futures.TimeoutError:
        logging.error(f"Processing timed out for args: {args}")
        if future:
            future.cancel()
        return None
    except Exception as e:
        logging.error(f"Error in process_with_timeout: {e}")
        return None
    
    return result

def run_subprocess_with_timeout(cmd, timeout=DEFAULT_PROCESS_TIMEOUT, **kwargs):
    """
    Run a subprocess command with timeout.
    
    Parameters:
        cmd (list): Command list to execute
        timeout (int): Timeout in seconds
        **kwargs: Additional arguments to pass to subprocess.run
        
    Returns:
        subprocess.CompletedProcess or None: Result of the process or None if timeout
    """
    try:
        return subprocess.run(cmd, timeout=timeout, **kwargs)
    except subprocess.TimeoutExpired:
        logging.error(f"Command {' '.join(cmd)} timed out after {timeout} seconds")
        return None
    except Exception as e:
        logging.error(f"Error running command {' '.join(cmd)}: {e}")
        return None

def get_installed_tesseract_langs():
    """
    Retrieve a list of installed Tesseract languages.
    """
    try:
        langs_output = subprocess.check_output(['tesseract', '--list-langs'], stderr=subprocess.STDOUT).decode('utf-8')
        langs = langs_output.splitlines()
        if len(langs) > 1:
            return langs[1:]
        return []
    except Exception as e:
        logging.error(f"Error retrieving installed Tesseract languages: {e}")
        return []

# Function to check and install required Tesseract languages
def check_and_install_tesseract_langs(required_langs):
    """
    Check if the required Tesseract language data files are present.
    If not, attempt to find tessdata directories and prompt the user to install missing languages.
    """
    tessdata_dirs = find_tessdata_dirs()

    if not tessdata_dirs:
        logging.warning("Could not find 'tessdata' directory. Please ensure Tesseract OCR is correctly installed.")
        return False

    # Check for each required language
    missing_langs = []
    for lang in required_langs:
        lang_file = f"{lang}.traineddata"
        found = False
        for tessdata_dir in tessdata_dirs:
            if os.path.isfile(os.path.join(tessdata_dir, lang_file)):
                found = True
                break
        if not found:
            missing_langs.append(lang)

    if not missing_langs:
        logging.debug("All required Tesseract language data files are present.")
        return True

    logging.warning(f"Missing Tesseract language data files: {', '.join(missing_langs)}")

    # Determine OS
    current_os = platform.system()
    install_cmd = None
    if current_os == "Darwin":  # macOS
        logging.info("Detected macOS.")
        if shutil.which('brew') is None:
            logging.error("Homebrew is not installed. Please install Homebrew to proceed.")
            return False
        # Construct brew install commands for missing languages
        brew_langs = [f"tesseract-lang-{lang}" for lang in missing_langs]
        install_cmd = ['brew', 'install'] + brew_langs
    elif current_os == "Linux":
        logging.info("Detected Linux.")
        # Attempt to detect package manager
        if shutil.which('apt-get'):
            install_cmd = ['sudo', 'apt-get', 'install', '-y'] + [f"tesseract-ocr-{lang}" for lang in missing_langs]
        elif shutil.which('yum'):
            install_cmd = ['sudo', 'yum', 'install', '-y'] + [f"tesseract-langpack-{lang}" for lang in missing_langs]
        else:
            logging.error("Unsupported package manager. Please install the required Tesseract language data manually.")
            return False
    else:
        logging.error(f"Unsupported operating system: {current_os}. Please install the required Tesseract language data manually.")
        return False

    # Prompt user for installation
    user_input = 'n'
    try:
        user_input = input(f"Do you want to install the missing Tesseract language data ({', '.join(missing_langs)})? [y/N]: ").strip().lower()
    except EOFError:
        logging.warning("Input not available. Skipping installation of Tesseract languages.")
        return False

    if user_input == 'y':
        try:
            logging.info(f"Installing Tesseract language data: {' '.join(missing_langs)}")
            subprocess.run(install_cmd, check=True)
            logging.info("Tesseract language data installation completed.")
            return True
        except subprocess.CalledProcessError as e:
            logging.error(f"Error installing Tesseract language data: {e}")
        except Exception as e:
            logging.error(f"Unexpected error during installation: {e}")
    else:
        logging.info("User opted not to install the missing Tesseract language data. OCR functionality may be limited.")
    return False

def ensure_kraken_models(models, verbose=False):
    """
    Ensure that all required Kraken models are downloaded.

    Parameters:
        models (list): List of model IDs to check and download.
        verbose (bool): Enable verbose logging.

    Returns:
        None
    """
    for model_id in models:
        try:
            # Check if the model is already downloaded by attempting to show it
            result = subprocess.run(['kraken', 'show', model_id], capture_output=True, text=True)
            if result.returncode == 0:
                logging.debug(f"Model {model_id} is already downloaded.")
            else:
                logging.info(f"Model {model_id} not found. Downloading...")
                subprocess.run(['kraken', 'get', model_id], check=True)
                logging.info(f"Successfully downloaded model: {model_id}")
        except subprocess.CalledProcessError as e:
            logging.error(f"Failed to ensure Kraken model {model_id}: {e}")
        except Exception as e:
            logging.error(f"Unexpected error while ensuring Kraken model {model_id}: {e}")

# Text extraction functions for various file types
def extract_text_with_pdfminer(pdf_path, verbose=False):
    if not optional_libraries.get('pdfminer', False):
        logging.warning(f"pdfminer is not available. Skipping PDFMiner extraction for {pdf_path}.")
        return ""
    logging.debug(f"Extracting text from PDF using pdfminer: {pdf_path}")
    from pdfminer.pdfparser import PDFSyntaxError
    try:
        text = pdfminer_extract_text(pdf_path, maxpages=4)
        logging.debug(f"Successfully extracted text using pdfminer ({len(text)} characters)")
        if verbose:
            logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
        return text[:2000]
    except PDFSyntaxError as e:
        logging.error(f"PDF syntax error in {pdf_path}: {e}")
        return ""
    except Exception as e:
        logging.error(f"Error extracting text from {pdf_path} with pdfminer: {e}")
        return ""


def perform_ocr_on_image(image, ocr_available=None, tools_available=None, optional_libraries=None, verbose=False, ocr_attempts=None):
    """
    Perform OCR on an image using available OCR tools in preferred order.
    
    Parameters:
        image (PIL.Image.Image): PIL Image object.
        ocr_available (dict): Availability of OCR tools.
        tools_available (dict): Dictionary indicating availability of tools.
        optional_libraries (dict): Dictionary indicating availability of optional libraries.
        verbose (bool): Enable verbose logging.
        ocr_attempts (set): Set to track which OCR tools have been attempted.
    
    Returns:
        str: Extracted text from OCR or empty string if OCR fails.
    """
    if ocr_available is None:
        ocr_available = {}
    if tools_available is None:
        tools_available = {}
    if optional_libraries is None:
        optional_libraries = {}
    if ocr_attempts is None:
        ocr_attempts = set()
    
    text = ""
    
    # Determine available OCR methods based on preferred order
    available_ocr_methods = [method for method in OCR_PREFERRED_ORDER if ocr_available.get(method, False) and method not in ocr_attempts]
    if not available_ocr_methods:
        logging.warning("No OCR tools are available or have not been attempted yet.")
        return text
    
    for ocr_method in available_ocr_methods:
        if shutdown_flag.is_set():
            logging.info("Shutdown initiated. Stopping OCR on image.")
            break
        
        if ocr_method == 'tesseract' and image and optional_libraries.get('pytesseract', False):
            # Make sure image is actually a PIL Image
            if not hasattr(image, 'mode'):
                logging.warning("Invalid image object for Tesseract OCR.")
                continue
                
            lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
            if not lang:
                lang = 'eng'  # Default to English if no languages are found
                
            logging.debug(f"Performing OCR on image with Tesseract using languages: {lang}")
            try:
                import pytesseract
                page_text = pytesseract.image_to_string(image, lang=lang)
                if page_text.strip() and is_text_meaningful(page_text, file_path="image"):
                    text += page_text + "\n"
                    logging.debug(f"Successfully extracted meaningful text using Tesseract ({len(page_text)} characters)")
                    if verbose:
                        logging.debug(f"Extracted text with Tesseract: {page_text[:500]}")
                    return text
                else:
                    logging.debug("Text extracted with Tesseract is not meaningful. Trying next OCR tool.")
            except Exception as e:
                logging.error(f"Error performing OCR with Tesseract on image: {e}")
                ocr_attempts.add('tesseract')
        
        elif ocr_method == 'kraken' and image:
            logging.debug("Performing OCR on image with Kraken.")
            page_text = perform_ocr_with_kraken(image, tools_available=tools_available, verbose=verbose, ocr_attempts=ocr_attempts)
            if page_text.strip() and is_text_meaningful(page_text, file_path="image"):
                text += page_text + "\n"
                logging.debug(f"Successfully extracted meaningful text using Kraken ({len(page_text)} characters)")
                if verbose:
                    logging.debug(f"Extracted text with Kraken: {page_text[:500]}")
                return text
            else:
                logging.debug("Text extracted with Kraken is not meaningful. Trying next OCR tool.")
        
        # Mark this OCR method as attempted
        ocr_attempts.add(ocr_method)
    
    # If we've tried all OCR methods and none worked, try using tesseract directly via subprocess
    if 'tesseract' not in ocr_attempts and shutil.which('tesseract') and hasattr(image, 'save'):
        logging.debug("Attempting OCR with Tesseract via subprocess...")
        try:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_path = temp_file.name
            
            image.save(temp_path)
            
            lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
            if not lang:
                lang = 'eng'  # Default to English if no languages are found
                
            cmd = ['tesseract', temp_path, 'stdout', '-l', lang]
            result = run_subprocess_with_timeout(cmd, timeout=OCR_PROCESS_TIMEOUT, capture_output=True, text=True)
            
            # Clean up the temporary file
            try:
                os.unlink(temp_path)
            except:
                pass
            
            if result and result.returncode == 0:  # Only process if successful
                page_text = result.stdout.strip()
                if page_text.strip() and is_text_meaningful(page_text, file_path="image"):
                    text += page_text + "\n"
                    logging.debug(f"Successfully extracted meaningful text using Tesseract subprocess ({len(page_text)} characters)")
                    if verbose:
                        logging.debug(f"Extracted text with Tesseract subprocess: {page_text[:500]}")
                    return text
        except Exception as e:
            logging.error(f"Error performing OCR with Tesseract subprocess: {e}")
        finally:
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except:
                    pass
    
    return text

def perform_ocr_with_kraken(image, model='en_best.mlmodel', tools_available=None, verbose=False, ocr_attempts=None):
    """
    Perform OCR using Kraken on a given image, with subprocess fallback.
    
    Parameters:
        image (PIL.Image.Image): PIL Image object.
        model (str): Path or name of the Kraken model to use.
        tools_available (dict): Dictionary indicating availability of tools.
        verbose (bool): Enable verbose logging.
        ocr_attempts (set): Set to track which OCR tools have been attempted.
    
    Returns:
        str: Extracted text from Kraken OCR or empty string if OCR fails.
    """
    if tools_available is None:
        tools_available = {}
    if ocr_attempts is None:
        ocr_attempts = set()
    
    # Global counter to track Kraken failures
    global kraken_failure_count
    if not hasattr(globals(), 'kraken_failure_count'):
        globals()['kraken_failure_count'] = 0
    
    # Check if we've had too many failures
    if globals()['kraken_failure_count'] >= 3:
        logging.warning("Too many Kraken failures. Skipping Kraken OCR for this session.")
        tools_available['kraken'] = False
        ocr_attempts.add('kraken')
        return ""
    
    if not tools_available.get('kraken', False) or 'kraken' not in shutil.which('kraken'):
        logging.warning("Kraken OCR is not available. Skipping Kraken OCR.")
        return ""
    
    # Try direct subprocess approach first (most reliable)
    try:
        logging.debug("Attempting OCR with Kraken via direct subprocess...")
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            temp_path = temp_file.name
        
        # Save the image to a temporary file
        image.save(temp_path)
        logging.debug(f"Saved image to temporary file {temp_path}")
        
        # Create a temporary output file
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as out_file:
            output_path = out_file.name
        
        try:
            # Run kraken OCR command
            cmd = ['kraken', '-i', temp_path, output_path, 'ocr', '-m', model]
            result = run_subprocess_with_timeout(cmd, timeout=OCR_PROCESS_TIMEOUT, capture_output=True, text=True)
            
            if result and result.returncode == 0:
                # Read the output file
                with open(output_path, 'r', encoding='utf-8') as f:
                    text = f.read().strip()
                
                if text and is_text_meaningful(text, file_path="image"):
                    logging.debug(f"Successfully extracted meaningful text using Kraken subprocess ({len(text)} characters)")
                    if verbose:
                        logging.debug(f"Extracted text with Kraken subprocess: {text[:500]}")
                    return text
            else:
                if result:
                    logging.error(f"Kraken subprocess failed with return code {result.returncode}: {result.stderr}")
                globals()['kraken_failure_count'] += 1
        except Exception as e:
            logging.error(f"Error running Kraken subprocess: {e}")
            globals()['kraken_failure_count'] += 1
        finally:
            # Clean up temporary files
            for path in [temp_path, output_path]:
                if os.path.exists(path):
                    try:
                        os.unlink(path)
                    except:
                        pass
    except Exception as e:
        logging.error(f"Error in Kraken subprocess preparation: {e}")
        globals()['kraken_failure_count'] += 1
    
    # If we've reached here, try the Python module approach as a fallback
    try:
        # Try to import the required modules at runtime
        from kraken import binarization
        from kraken import pageseg
        
        # Try the newer Kraken API (kraken 3.x+)
        try:
            # Try first with the fully qualified version
            from kraken.recognition import load_any as load_model
            from kraken.recognition import ocr_record
            
            if verbose:
                logging.debug("Using Kraken recognition API (newer version)")
            
            # Binarize the image
            bin_img = binarization.nlbin(image)
            
            # Segment the image into lines
            regions = pageseg.segment(bin_img)
            
            # Load the model
            model_obj = load_model(model)
            
            # Perform OCR
            result = ocr_record(network=model_obj, regions=regions, im=bin_img)
            
            # Extract text from result
            text = ''
            for line in result:
                text += line.prediction + '\n'
                
            if text and is_text_meaningful(text, file_path="image"):
                return text
            
        except (ImportError, AttributeError) as e:
            if verbose:
                logging.debug(f"Failed to use recognition API: {e}. Trying rpred API.")
                
            # Try with rpred (kraken 2.x)
            try:
                from kraken import rpred
                
                if verbose:
                    logging.debug("Using Kraken rpred API")
                
                # Binarize the image
                bin_img = binarization.nlbin(image)
                
                # Segment the image into lines
                regions = pageseg.segment(bin_img)
                
                # Perform OCR
                result = rpred.rpred(model, bin_img, regions)
                text = result.prediction if hasattr(result, 'prediction') else ''
                
                if text and is_text_meaningful(text, file_path="image"):
                    return text
                
            except (ImportError, AttributeError) as e:
                if verbose:
                    logging.debug(f"Failed to use rpred API: {e}")
                globals()['kraken_failure_count'] += 1
                
    except ImportError as e:
        logging.error(f"Kraken OCR dependencies are missing: {e}")
        globals()['kraken_failure_count'] += 1
        ocr_attempts.add('kraken')
    except Exception as e:
        logging.error(f"Error performing OCR with Kraken: {e}")
        globals()['kraken_failure_count'] += 1
        ocr_attempts.add('kraken')
    
    return ""

# ensure Tesseract is properly checked and available
def check_tesseract_availability():
    """Check if Tesseract is available and properly installed"""
    tesseract_available = False
    if shutil.which('tesseract'):
        try:
            # Verify that tesseract works by running a simple version command
            result = subprocess.run(['tesseract', '--version'], 
                                   capture_output=True, text=True, check=True)
            tesseract_available = True
            logging.debug(f"Tesseract is available: {result.stdout.splitlines()[0] if result.stdout else 'Unknown version'}")
            
            # Verify pytesseract can be imported if available
            try:
                import pytesseract
                optional_libraries['pytesseract'] = True
                logging.debug("pytesseract Python module is available")
            except ImportError:
                logging.warning("pytesseract Python module is not available. Will use subprocess for Tesseract OCR.")
                optional_libraries['pytesseract'] = False
        except subprocess.CalledProcessError as e:
            logging.error(f"Error verifying Tesseract: {e}")
            tesseract_available = False
    else:
        logging.warning("Tesseract is not installed or not in PATH")
        tesseract_available = False
        
    return tesseract_available


def extract_text_from_image(image_path, perform_ocr_if_needed, ocr_available, verbose=False, tools_available=None, optional_libraries=None, ocr_attempts=None):
    """
    Extract text from an image file using OCR methods first, then fall back to textract.
    
    Parameters:
        image_path (str): Path to the image file.
        perform_ocr_if_needed (bool): Whether to perform OCR if textract fails.
        ocr_available (dict): Availability of OCR tools.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
        optional_libraries (dict): Dictionary indicating availability of optional libraries.
    
    Returns:
        str: Extracted text or empty string if extraction fails.
    """
    text = ""
    if tools_available is None:
        tools_available = {}
    if optional_libraries is None:
        optional_libraries = {}

    # Proceed with OCR if needed
    if perform_ocr_if_needed and (ocr_available.get('tesseract', False) or ocr_available.get('kraken', False)):
        logging.debug("Attempting OCR on image.")

        image = None
        if tools_available.get('pillow', False):
            try:
                image = Image.open(image_path)
            except Exception as e:
                logging.error(f"Error opening image {image_path} for OCR: {e}")
                image = None

        if image:
            text = perform_ocr_on_image(
                image,
                ocr_available,
                tools_available=tools_available,
                optional_libraries=optional_libraries,
                verbose=verbose,
                ocr_attempts=ocr_attempts,
            )
        else:
            # Fall back to direct Tesseract subprocess if Pillow is not available or image failed to open
            logging.debug("Pillow is not available or failed to open image. Performing OCR via Tesseract subprocess.")
            try:
                lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
                cmd = ['tesseract', image_path, 'stdout', '-l', lang]
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                text = result.stdout.strip()
                logging.debug(f"Successfully extracted text using Tesseract subprocess ({len(text)} characters)")
                if verbose:
                    logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
            except subprocess.CalledProcessError as e:
                logging.error(f"Error performing OCR with Tesseract subprocess on {image_path}: {e}")
            except Exception as e:
                logging.error(f"Unexpected error during OCR with Tesseract subprocess on {image_path}: {e}")

        # Validate the extracted text
        if is_text_meaningful(text, file_path=image_path):
            logging.info(f"Extracted meaningful text from {image_path} using OCR.")
            return text[:2000]
        else:
            logging.warning(f"OCR extraction from {image_path} did not yield meaningful text.")

    # Attempt text extraction with textract as a fallback
    if tools_available.get('textract', False):
        logging.debug(f"Attempting text extraction from image using textract: {image_path}")
        try:
            result = textract.process(image_path)
            if result:
                text = result.decode('utf-8').strip()
                if text:
                    logging.debug(f"Successfully extracted text using textract ({len(text)} characters)")
                    if verbose:
                        logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                    logging.info(f"Extracted text from {image_path} using textract.")
                    return text[:2000]
        except Exception as e:
            logging.error(f"Error extracting text from {image_path} using textract: {e}")

    logging.warning(f"Failed to extract meaningful text from image: {image_path}.")
    return ""


def extract_text_from_docx(docx_path, verbose=False, tools_available=None):
    """
    Extract text from a DOCX file using python-docx.
    Falls back to textract if python-docx is unavailable or fails.
    
    Parameters:
        docx_path (str): Path to the DOCX file.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
    
    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if tools_available is None:
        tools_available = {}
    
    if not tools_available.get('docx', False):
        logging.warning(f"DOCX processing is skipped for {docx_path} as python-docx is not available.")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(docx_path, verbose=verbose)
        else:
            logging.error("Textract is also unavailable. Cannot extract text from DOCX.")
            return ""
    
    logging.debug(f"Extracting text from DOCX: {docx_path}")
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
            if len(text) > 3000:
                break
        logging.debug(f"Extracted {len(text)} characters from DOCX")
        if verbose:
            logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
    except Exception as e:
        logging.error(f"Error extracting text from {docx_path}: {e}")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            text = extract_text_with_textract(docx_path, verbose=verbose)
    return text[:2000]


def is_text_meaningful(text, min_length=200, min_unique_words=20, alnum_ratio_threshold=0.3, max_repetition_ratio=0.3, file_path=""):
    if len(text) < min_length:
        logging.debug(f"{file_path} - Text length {len(text)} is below the minimum threshold of {min_length}.")
        return False

    alnum_count = sum(c.isalnum() for c in text)
    alnum_ratio = alnum_count / len(text)
    if alnum_ratio < alnum_ratio_threshold:
        logging.debug(f"{file_path} - Alphanumeric ratio {alnum_ratio:.2f} is below the threshold of {alnum_ratio_threshold}.")
        return False

    words = re.findall(r'\b\w+\b', text.lower())
    if not words:
        logging.debug(f"{file_path} - No words found in text.")
        return False

    unique_words = set(words)
    if len(unique_words) < min_unique_words:
        logging.debug(f"{file_path} - Unique words count {len(unique_words)} is below the threshold of {min_unique_words}.")
        return False

    word_counts = {}
    for word in words:
        word_counts[word] = word_counts.get(word, 0) + 1

    max_repetition = max(word_counts.values()) / len(words)
    if max_repetition > max_repetition_ratio:
        logging.debug(f"{file_path} - Max word repetition ratio {max_repetition:.2f} exceeds the threshold of {max_repetition_ratio}.")
        return False

    return True


def extract_text_from_mobi(mobi_path, verbose=False, tools_available=None):
    """
    Extract text from a MOBI file using the mobi library.
    Falls back to textract if mobi is unavailable or fails.
    
    Parameters:
        mobi_path (str): Path to the MOBI file.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
    
    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if tools_available is None:
        tools_available = {}
    
    if not tools_available.get('mobi', False):
        logging.warning(f"MOBI processing is skipped for {mobi_path} as mobi library is not available.")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(mobi_path, verbose=verbose)
        else:
            logging.error("Textract is also unavailable. Cannot extract text from MOBI.")
            return ""
    
    logging.debug(f"Extracting text from MOBI: {mobi_path}")
    text = ""
    try:
        with tempfile.TemporaryDirectory() as tempdir:
            filepath = mobi.extract(mobi_path, tempdir=tempdir)
            logging.debug(f"mobi.extract returned filepath: {filepath}")
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
                logging.debug(f"Successfully extracted text using mobi library ({len(text)} characters)")
                if verbose:
                    logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
    except Exception as e:
        logging.error(f"Error extracting text from {mobi_path} using mobi library: {e}")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(mobi_path, verbose=verbose)
    return text[:2000]


def extract_text_from_djvu(djvu_path, djvutxt_installed, verbose=False, tools_available=None):
    """
    Extract text from a DJVU file using djvutxt.
    Falls back to textract if djvutxt is unavailable or fails.
    
    Parameters:
        djvu_path (str): Path to the DJVU file.
        djvutxt_installed (bool): Whether djvutxt is installed.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
    
    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if tools_available is None:
        tools_available = {}
    
    if not djvutxt_installed:
        logging.warning(f"Skipping DJVU file {djvu_path} as 'djvutxt' is not available.")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(djvu_path, verbose=verbose)
        else:
            logging.error("Textract is also unavailable. Cannot extract text from DJVU.")
            return ""
    
    logging.debug(f"Extracting text from DJVU: {djvu_path}")
    text = ""
    try:
        result = subprocess.run(['djvutxt', djvu_path], capture_output=True, text=True, check=True)
        text = result.stdout[:2000]
        logging.debug(f"Successfully extracted {len(text)} characters from DJVU")
        if verbose:
            logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
    except subprocess.CalledProcessError as e:
        logging.error(f"Error extracting text from {djvu_path}: {e}")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(djvu_path, verbose=verbose)
    except FileNotFoundError:
        logging.error("djvutxt command not found. Please install DjVuLibre package.")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(djvu_path, verbose=verbose)
    return text[:2000]


def extract_text_from_azw(azw_path, verbose=False, tools_available=None):
    """
    Extract text from an AZW file using mobi library and EbookLib.
    Falls back to textract if necessary.
    
    Parameters:
        azw_path (str): Path to the AZW file.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
    
    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if tools_available is None:
        tools_available = {}
    
    if not tools_available.get('mobi', False):
        logging.warning(f"AZW processing is skipped for {azw_path} as mobi library is not available.")
        if tools_available.get('textract', False):
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(azw_path, verbose=verbose)
        else:
            logging.error("Textract is also unavailable. Cannot extract text from AZW.")
            return ""
    
    logging.debug(f"Extracting text from AZW: {azw_path}")
    text = ""
    if tools_available.get('textract', False):
        logging.debug("Attempting to use textract for AZW extraction...")
        try:
            text = textract.process(azw_path).decode('utf-8')
            if text.strip():
                logging.debug(f"Successfully extracted text using textract ({len(text)} characters)")
                if verbose:
                    logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                logging.info(f"Extracted text from {azw_path} using textract.")
                return text[:2000]
        except Exception as e:
            logging.error(f"Error extracting text from {azw_path} using textract: {e}")
    
    if tools_available.get('ebooklib', False):
        logging.debug("Attempting to use EbookLib for AZW extraction...")
        try:
            book = epub.read_epub(azw_path)
            for item in book.get_items():
                if item.get_type() == EPUB_ITEM_DOCUMENT:
                    if tools_available.get('beautifulsoup4', False):
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        text += soup.get_text() + "\n"
                    else:
                        logging.warning("beautifulsoup4 is not available. Skipping HTML parsing.")
                    if len(text) > 3000:
                        break
            if text.strip():
                logging.debug(f"Successfully extracted text using EbookLib ({len(text)} characters)")
                if verbose:
                    logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                logging.info(f"Extracted text from {azw_path} using EbookLib.")
                return text[:2000]
        except Exception as e:
            logging.error(f"Error extracting text from {azw_path} using EbookLib: {e}")
    
    if tools_available.get('mobi', False):
        logging.debug("Attempting to use mobi library for AZW extraction...")
        tempdir = None  # Ensure tempdir is initialized
        try:
            tempdir, filepath = mobi.extract(azw_path)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
                    logging.debug(f"Successfully extracted text using mobi library ({len(text)} characters)")
                    if verbose:
                        logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
            finally:
                if tempdir:
                    shutil.rmtree(tempdir)
            if text.strip():
                return text[:2000]
            else:
                logging.warning(f"No meaningful text extracted from {azw_path} using mobi library.")
                if tools_available.get('textract', False):
                    logging.info("Attempting to extract text using textract as a fallback.")
                    return extract_text_with_textract(azw_path, verbose=verbose)
        except Exception as e:
            logging.error(f"Error extracting text from {azw_path} using mobi library: {e}")
            if tempdir:
                try:
                    shutil.rmtree(tempdir)
                except Exception as cleanup_e:
                    logging.error(f"Error removing temporary directory {tempdir}: {cleanup_e}")
            if tools_available.get('textract', False):
                logging.info("Attempting to extract text using textract as a fallback.")
                return extract_text_with_textract(azw_path, verbose=verbose)
    return text[:2000]


def extract_text_from_pdf(pdf_path, perform_ocr_if_needed, ocr_available, verbose=False, tools_available=None, ocr_attempts=None):
    """
    Extract text from a PDF file using various methods in priority order.
    
    Parameters:
        pdf_path (str): Path to the PDF file.
        perform_ocr_if_needed (bool): Whether to perform OCR if text extraction fails.
        ocr_available (dict): Availability of OCR tools.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
        ocr_attempts (set): Set to track which OCR tools have been attempted.
    
    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if tools_available is None:
        tools_available = {}
    
    text = ""

    # Check if the PDF is valid
    if not is_pdf_valid(pdf_path, tools_available=tools_available):
        logging.error(f"Invalid PDF file: {pdf_path}. Skipping extraction.")
        return ""

    # Method 1: Try extracting with PyMuPDF (fitz)
    if tools_available.get('fitz', False):
        logging.debug(f"Attempting text extraction from PDF using PyMuPDF (fitz): {pdf_path}")
        try:
            import fitz
            doc = fitz.open(pdf_path)
            for page_num in range(min(6, doc.page_count)):  # Limit to first 6 pages
                page = doc[page_num]
                extracted_text = page.get_text()
                if extracted_text:
                    text += extracted_text + "\n"
                if len(text) > 3000:
                    break
            doc.close()
            if text.strip():
                logging.debug(f"Successfully extracted text using PyMuPDF ({len(text)} characters)")
                if verbose:
                    logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                logging.info(f"Extracted meaningful text from {pdf_path} using PyMuPDF.")
                return text[:2000]
        except Exception as e:
            logging.error(f"Error extracting text from {pdf_path} using PyMuPDF: {e}")

    # Method 2: Try extracting with pdfplumber
    if tools_available.get('pdfplumber', False):
        logging.debug(f"Attempting text extraction from PDF using pdfplumber: {pdf_path}")
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages[:6]:  # Limit to first 6 pages
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + "\n"
                    if len(text) > 3000:
                        break
            if text.strip():
                logging.debug(f"Successfully extracted text using pdfplumber ({len(text)} characters)")
                if verbose:
                    logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                logging.info(f"Extracted meaningful text from {pdf_path} using pdfplumber.")
                return text[:2000]
        except Exception as e:
            logging.error(f"Error extracting text from {pdf_path} using pdfplumber: {e}")

    # Method 3: Try extracting with PyPDF2
    if tools_available.get('PyPDF2', False):
        logging.debug(f"Attempting to extract text from PDF using PyPDF2: {pdf_path}")
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(pdf_path)
            for page_num in range(min(6, len(reader.pages))):
                logging.debug(f"Reading page: {page_num}")
                page = reader.pages[page_num]
                extracted_text = page.extract_text() or ''
                if extracted_text:
                    text += extracted_text + "\n"
                if len(text) > 3000:
                    break
            if is_text_meaningful(text, file_path=pdf_path):
                logging.debug(f"Successfully extracted text using PyPDF2 ({len(text)} characters)")
                if verbose:
                    logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                logging.info(f"Extracted meaningful text from {pdf_path} using PyPDF2.")
                return text[:2000]
            else:
                logging.warning(f"{pdf_path} - Text extracted by PyPDF2 is not meaningful.")
        except (PdfReadError, ValueError, TypeError) as e:
            logging.error(f"Error extracting text from {pdf_path} with PyPDF2: {e}.")

    # Method 4: Try extracting with pdfminer
    if tools_available.get('pdfminer', False):
        logging.debug(f"Attempting text extraction from PDF using pdfminer: {pdf_path}")
        try:
            text = extract_text_with_pdfminer(pdf_path, verbose=verbose)
            if is_text_meaningful(text, file_path=pdf_path):
                logging.info(f"Extracted meaningful text from {pdf_path} using pdfminer.")
                return text[:2000]
            else:
                logging.warning(f"{pdf_path} - Text extracted by pdfminer is not meaningful.")
        except Exception as e:
            logging.error(f"Error extracting text from {pdf_path} using pdfminer: {e}")

    # Method 5: Try OCR (Tesseract or Kraken)
    if perform_ocr_if_needed and (ocr_available.get('tesseract', False) or ocr_available.get('kraken', False)):
        logging.debug("Attempting OCR due to failure of text extraction.")
        text = perform_ocr_on_first_pages(
            pdf_path,
            num_pages=6,
            verbose=verbose,
            tools_available=tools_available,
            ocr_attempts=ocr_attempts
        )
        if text.strip():
            logging.debug(f"Successfully extracted text using OCR ({len(text)} characters)")
            if verbose:
                logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
            logging.info(f"Successfully performed OCR on {pdf_path}.")
            return text[:2000]
        else:
            logging.warning(f"OCR extraction from {pdf_path} did not yield meaningful text.")

    # Method 6: Last resort - Try extracting with textract using pdftotext
    if tools_available.get('textract', False):
        logging.debug(f"Attempting text extraction from PDF using textract with pdftotext: {pdf_path}")
        try:
            # Explicitly use 'pdftotext' as the method to avoid Python 2 dependency
            result = textract.process(pdf_path, method='pdftotext')
            if result:
                text = result.decode('utf-8').strip()
                if text:
                    logging.debug(f"Successfully extracted text using textract with pdftotext ({len(text)} characters)")
                    if verbose:
                        logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                    logging.info(f"Extracted text from {pdf_path} using textract with pdftotext.")
                    return text[:2000]
        except Exception as e:
            logging.error(f"Error extracting text from {pdf_path} using textract with pdftotext: {e}")

    logging.warning(f"Failed to extract meaningful text from PDF: {pdf_path}")
    return ""


def is_pdf_valid(file_path, tools_available=None, optional_libraries=None):
    """
    Validate the integrity of a PDF file using available tools.
    Utilizes PyPDF2, pdfminer, and textract with graceful fallbacks.
    
    Parameters:
        file_path (str): Path to the PDF file.
        tools_available (dict): Dictionary indicating availability of tools.
        optional_libraries (dict): Dictionary indicating availability of optional libraries.
    
    Returns:
        bool: True if PDF is valid, False otherwise.
    """
    if tools_available is None:
        tools_available = {}
    if optional_libraries is None:
        optional_libraries = {}
    
    logging.debug(f"Validating PDF file: {file_path}")
    validation_methods = []
    
    if optional_libraries.get('PyPDF2', False):
        validation_methods.append('PyPDF2')
    if optional_libraries.get('pdfminer', False):
        validation_methods.append('pdfminer')
    if tools_available.get('textract', False):
        validation_methods.append('textract')
    
    for method in validation_methods:
        if shutdown_flag.is_set():
            logging.info(f"Shutdown initiated. Stopping PDF validation for {file_path}.")
            return False
        logging.debug(f"Attempting PDF validation using {method} for {file_path}")
        if method == 'PyPDF2':
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    num_pages = len(reader.pages)
                    logging.debug(f"PyPDF2 successfully read {num_pages} pages.")
                    return True
            except Exception as e:
                logging.error(f"PyPDF2 failed to read {file_path}: {e}")
        elif method == 'pdfminer':
            try:
                text = pdfminer_extract_text(file_path, maxpages=1)
                if text.strip():
                    logging.debug(f"pdfminer successfully extracted text from {file_path}.")
                    return True
                else:
                    logging.error(f"pdfminer extracted no text from {file_path}.")
            except Exception as e:
                logging.error(f"pdfminer failed to extract text from {file_path}: {e}")
        elif method == 'textract':
            try:
                text = textract.process(file_path, method='pdfminer').decode('utf-8')
                if text.strip():
                    logging.debug(f"textract successfully extracted text from {file_path}.")
                    return True
                else:
                    logging.error(f"textract extracted no text from {file_path}.")
            except Exception as e:
                logging.error(f"textract failed to extract text from {file_path}: {e}")
    
    logging.warning(f"All PDF validation methods failed for {file_path}. Marking as invalid.")
    return False

def perform_ocr_on_first_pages(pdf_path, num_pages=3, verbose=False, tools_available=None, language='en', ocr_attempts=None):
    """
    Perform OCR on the first few pages of a PDF file using available OCR tools in preferred order.
    
    Parameters:
        pdf_path (str): Path to the PDF file.
        num_pages (int): Number of pages to perform OCR on.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
        language (str): Language code for selecting the Kraken model.
        ocr_attempts (set): Set to track which OCR tools have been attempted.
    
    Returns:
        str: Extracted text up to 2000 characters or empty string if OCR fails.
    """
    if tools_available is None:
        tools_available = {}
    if ocr_attempts is None:
        ocr_attempts = set()
    
    if not (tools_available.get('pdf2image', False)):
        logging.warning("pdf2image is not available. OCR on specific pages will be skipped.")
        return ""
    logging.debug(f"Performing OCR on the first {num_pages} pages of: {pdf_path}")
    text = ""
    try:
        # Dynamically find tessdata directories and set TESSDATA_PREFIX
        tessdata_dirs = find_tessdata_dirs()
        if tessdata_dirs:
            TESSDATA_PREFIX = tessdata_dirs[0]  # Use the first found tessdata directory
            os.environ['TESSDATA_PREFIX'] = TESSDATA_PREFIX
            logging.debug(f"Set TESSDATA_PREFIX to: {TESSDATA_PREFIX}")
        else:
            logging.warning("Could not find 'tessdata' directory. Please set the TESSDATA_PREFIX environment variable manually.")
            TESSDATA_PREFIX = None  # Handle cases where tessdata is not found

        # Try to directly convert the PDF pages to images
        try:
            images = convert_from_path(pdf_path, dpi=300, first_page=1, last_page=num_pages)
            logging.debug(f"Successfully converted {len(images)} pages from PDF to images")
        except Exception as e:
            logging.error(f"Error converting PDF to images: {e}")
            return ""

        # Try OCR on each image
        for i, image in enumerate(images):
            if shutdown_flag.is_set():
                logging.info(f"Shutdown initiated. Stopping OCR on {pdf_path}.")
                return text[:2000]
                
            # Try Tesseract directly - this is the most reliable method
            if shutil.which('tesseract'):
                try:
                    # Save the image to a temporary file
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                        temp_path = temp_file.name
                    
                    image.save(temp_path)
                    logging.debug(f"Saved page {i+1} to temporary file {temp_path}")
                    
                    # Get available languages
                    lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
                    if not lang:
                        lang = 'eng'  # Default to English if no required languages are found
                    
                    logging.debug(f"Performing OCR on page {i+1} with Tesseract subprocess using languages: {lang}")
                    cmd = ['tesseract', temp_path, 'stdout', '-l', lang]
                    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                    page_text = result.stdout.strip()
                    
                    # Clean up the temporary file
                    try:
                        os.unlink(temp_path)
                    except:
                        pass
                    
                    if page_text.strip() and is_text_meaningful(page_text, file_path=pdf_path):
                        text += page_text + "\n"
                        logging.debug(f"Successfully extracted meaningful text from page {i+1} using Tesseract subprocess ({len(page_text)} characters)")
                        if verbose:
                            logging.debug(f"Extracted Text from page {i+1} (first 500 chars): {page_text[:500]}")
                    else:
                        logging.debug(f"OCR with Tesseract on page {i+1} did not yield meaningful text.")
                except Exception as e:
                    logging.error(f"Error performing OCR with Tesseract subprocess on page {i+1}: {e}")
            
            # If Tesseract didn't work or isn't available, try with pytesseract Python library
            if not text.strip() and tools_available.get('pytesseract', False):
                try:
                    import pytesseract
                    lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
                    if not lang:
                        lang = 'eng'
                    logging.debug(f"Performing OCR on page {i+1} with pytesseract using languages: {lang}")
                    page_text = pytesseract.image_to_string(image, lang=lang)
                    if page_text.strip() and is_text_meaningful(page_text, file_path=pdf_path):
                        text += page_text + "\n"
                        logging.debug(f"Successfully extracted meaningful text from page {i+1} using pytesseract ({len(page_text)} characters)")
                        if verbose:
                            logging.debug(f"Extracted Text from page {i+1} (first 500 chars): {page_text[:500]}")
                    else:
                        logging.debug(f"OCR with pytesseract on page {i+1} did not yield meaningful text.")
                except Exception as e:
                    logging.error(f"Error performing OCR with pytesseract on page {i+1}: {e}")

            # If neither Tesseract method worked, try with Kraken
            if not text.strip() and tools_available.get('kraken', False):
                try:
                    from kraken import binarization, pageseg
                    logging.debug(f"Performing OCR on page {i+1} with Kraken.")
                    kraken_text = perform_ocr_with_kraken(image, model=get_kraken_model(language), tools_available=tools_available, verbose=verbose, ocr_attempts=ocr_attempts)
                    if kraken_text.strip() and is_text_meaningful(kraken_text, file_path=pdf_path):
                        text += kraken_text + "\n"
                        logging.debug(f"Successfully extracted meaningful text from page {i+1} using Kraken ({len(kraken_text)} characters)")
                        if verbose:
                            logging.debug(f"Extracted Text from page {i+1} (first 500 chars): {kraken_text[:500]}")
                    else:
                        logging.debug(f"OCR with Kraken on page {i+1} did not yield meaningful text.")
                except ImportError as e:
                    logging.error(f"Kraken OCR dependencies are missing for page {i+1}: {e}")
                except Exception as e:
                    logging.error(f"Error performing OCR with Kraken on page {i+1}: {e}")

            # Break if we've extracted enough text
            if len(text) > 2000:
                logging.debug(f"Extracted sufficient text (over 2000 chars). Stopping OCR.")
                break

        if text.strip():
            logging.debug(f"Successfully extracted text through OCR on first {num_pages} pages ({len(text)} characters)")
            if verbose:
                logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
            logging.info(f"Successfully performed OCR on first {num_pages} pages of {pdf_path}.")
        else:
            logging.warning(f"OCR failed to extract meaningful text from {pdf_path}")
    except Exception as e:
        logging.error(f"Error performing OCR on {pdf_path}: {e}")
    
    return text[:2000]


def extract_text_with_textract(file_path, verbose=False):
    if not essential_libraries.get('textract', False):
        logging.warning(f"Textract is not available. Skipping extraction for {file_path}.")
        return ""
    logging.debug(f"Extracting text with textract: {file_path}")
    try:
        if file_path.lower().endswith('.pdf'):
            text_bytes = textract.process(file_path, method='pdftotext')
        else:
            # Use default method for other file types
            text_bytes = textract.process(file_path)

        if text_bytes:
            result = text_bytes.decode('utf-8').strip()

        if result is None:
            logging.error(f"Textract returned None for {file_path}.")
            return ""
        if not isinstance(result, bytes):
            logging.error(f"Unexpected return type from textract for {file_path}: {type(result)}")
            return ""
        text = result.decode('utf-8')
        logging.debug(f"Successfully extracted {len(text)} characters with textract")
        if verbose:
            logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
        logging.info(f"Extracted text from {file_path} using textract.")
        return text[:2000]
    except Exception as e:
        logging.error(f"Error extracting text from {file_path} with textract: {e}")
        return ""

def extract_text_from_epub(epub_path, textract_available, tools_available, verbose=False):
    """
    Extract text from an EPUB file using EbookLib and BeautifulSoup4.
    Falls back to textract if EbookLib or BeautifulSoup4 is unavailable or fails.
    
    Parameters:
        epub_path (str): Path to the EPUB file.
        textract_available (bool): Whether textract is available for fallback.
        tools_available (dict): Dictionary indicating availability of tools.
        verbose (bool): Enable verbose logging.
    
    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if not tools_available.get('ebooklib', False):
        logging.warning(f"EPUB processing is skipped for {epub_path} as ebooklib is not available.")
        if textract_available:
            logging.info("Attempting to extract text using textract as a fallback.")
            return extract_text_with_textract(epub_path, verbose=verbose)
        else:
            logging.error("Textract is also unavailable. Cannot extract text from EPUB.")
            return ""
    
    logging.debug(f"Extracting text from EPUB: {epub_path}")
    text = ""
    try:
        book = epub.read_epub(epub_path)
        for item in book.get_items():
            if item.get_type() == EPUB_ITEM_DOCUMENT:
                if tools_available.get('beautifulsoup4', False):
                    try:
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        text += soup.get_text(separator="\n")
                    except Exception as e:
                        logging.error(f"Error parsing HTML content in {epub_path}: {e}")
                else:
                    logging.warning("beautifulsoup4 is not available. Skipping HTML parsing for this item.")
                if len(text) > 3000:
                    break
        if text.strip():
            logging.debug(f"Successfully extracted {len(text)} characters from EPUB using ebooklib.")
            if verbose:
                logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
            logging.info(f"Extracted text from {epub_path} using ebooklib.")
            return text[:2000]
        else:
            logging.warning(f"No meaningful text extracted from {epub_path} using ebooklib.")
            if textract_available:
                logging.info("Attempting to extract text using textract as a fallback.")
                return extract_text_with_textract(epub_path, verbose=verbose)
    except Exception as e:
        logging.error(f"Error extracting text from {epub_path} with ebooklib: {e}")
        if textract_available:
            logging.info("Attempting to extract text using textract as a fallback.")
            try:
                text = extract_text_with_textract(epub_path, verbose=verbose)
                if text.strip():
                    logging.debug(f"Successfully extracted text using textract from EPUB ({len(text)} characters)")
                    if verbose:
                        logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
                    logging.info(f"Extracted text from {epub_path} using textract.")
                    return text[:2000]
            except Exception as inner_e:
                logging.error(f"Error extracting text from {epub_path} using textract: {inner_e}")
    return ""

def extract_text_with_fitz(pdf_path, verbose=False):
    """
    Extract text from a PDF using PyMuPDF (fitz).

    Parameters:
        pdf_path (str): Path to the PDF file.
        verbose (bool): Enable verbose logging.

    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if not optional_libraries.get('fitz', False):
        logging.warning(f"PyMuPDF (fitz) is not available. Skipping PyMuPDF extraction for {pdf_path}.")
        return ""

    try:
        import fitz
        doc = fitz.open(pdf_path)
        text = ""
        for page_num in range(min(6, doc.page_count)):  # Limit to first 6 pages
            page = doc[page_num]
            extracted_text = page.get_text()
            if extracted_text:
                text += extracted_text + "\n"
            if len(text) > 3000:
                break
        doc.close()
        if text.strip():
            logging.debug(f"Successfully extracted text using PyMuPDF ({len(text)} characters)")
            if verbose:
                logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
            logging.info(f"Extracted meaningful text from {pdf_path} using PyMuPDF.")
            return text[:2000]
    except Exception as e:
        logging.error(f"Error extracting text from {pdf_path} using PyMuPDF: {e}")
    return ""

def extract_text_with_pdfplumber(pdf_path, verbose=False):
    """
    Extract text from a PDF using pdfplumber.

    Parameters:
        pdf_path (str): Path to the PDF file.
        verbose (bool): Enable verbose logging.

    Returns:
        str: Extracted text up to 2000 characters or empty string if extraction fails.
    """
    if not optional_libraries.get('pdfplumber', False):
        logging.warning(f"pdfplumber is not available. Skipping pdfplumber extraction for {pdf_path}.")
        return ""

    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:6]:  # Limit to first 6 pages
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
                if len(text) > 3000:
                    break
        if text.strip():
            logging.debug(f"Successfully extracted text using pdfplumber ({len(text)} characters)")
            if verbose:
                logging.debug(f"Extracted Text (first 500 chars): {text[:500]}")
            logging.info(f"Extracted meaningful text from {pdf_path} using pdfplumber.")
            return text[:2000]
    except Exception as e:
        logging.error(f"Error extracting text from {pdf_path} using pdfplumber: {e}")
    return ""


# Metadata handling functions
def clean_author_name(author_name):
    """Remove titles and punctuations from the author name."""
    author_name = re.sub(r'\bDr\.?\b', '', author_name, flags=re.IGNORECASE)  # Remove titles like Dr, Dr.
    author_name = re.sub(r'\s*,\s*', ' ', author_name).strip()  # Remove extra commas and spaces
    return author_name

def valid_author_name(author_name):
    """Check if the author name is valid, i.e., contains at least a first name and a last name,
    and does not contain special characters other than a-z, A-Z, hyphens, apostrophes, and maybe '.'.
    Also, ensure the name does not contain placeholders like 'lastname' or 'surname'."""
    parts = author_name.strip().split()
    if len(parts) <= 1:
        return False
    if not re.match(r'^[\w\s.\'-]+$', author_name, re.UNICODE):
        return False
    if "lastname" in author_name.lower() or "surname" in author_name.lower():
        return False
    return True

def sort_author_names(author_names, openai_client, max_attempts=5, verbose=False):
    for attempt in range(1, max_attempts + 1):
        if verbose:
            logging.debug(f"Attempt {attempt} to sort author names: {author_names}")
        formatted_author_names = author_names.replace('&', ',')
        prompt = (
            f"You will be given an author name that you must put into the format 'Lastname Surname'. "
            f"So, you must first make an educated guess if the given input is already in this format. If so, return it back. "
            f"If not and it is more plausibly in the format 'Surname(s) Lastname', you must reformat it. "
            f"Example: 'Michael Max Mustermann' must become 'Mustermann Michael Max' and 'Joe A. Doe' must become 'Doe Joe A'. "
            f"No comma after the Lastname! "
            f"If you are given multiple person names, only keep the first and omit all others. "
            f"If it is impossible to come up with a correct name, return <AUTHOR>n a</AUTHOR>. "
            f"You must give the output in the format: <AUTHOR>Lastname Surname(s)</AUTHOR>. "
            f"Here are the name parts: <AUTHOR>{formatted_author_names}</AUTHOR>"
        )
        messages = [{"role": "user", "content": prompt}]
        
        try:
            response = openai_client.chat.completions.create(
                model=MODEL_NAME,
                temperature=0.5,
                max_tokens=500,
                messages=messages
            )
        except Exception as e:
            logging.error(f"Error querying Ollama server for author names: {e}")
            return "UnknownAuthor"
        reformatted_name = response.choices[0].message.content.strip()
        logging.debug(f"Response from Ollama server: {reformatted_name}")
        name_match = re.search(r'<AUTHOR>(.*?)</AUTHOR>', reformatted_name)
        if name_match:
            ordered_name = name_match.group(1).strip().split(",")[0].strip()
            ordered_name = clean_author_name(ordered_name)
            logging.debug(f"Ordered name after cleaning: '{ordered_name}'")
            if valid_author_name(ordered_name):
                return ordered_name
            else:
                logging.warning(f"Invalid author name format detected: '{ordered_name}', retrying...")
        else:
            logging.warning("Failed to extract a valid name, retrying...")
    logging.error("Maximum retry attempts reached for sorting author names.")
    return "UnknownAuthor"


def parse_metadata(content, verbose=False):
    """
    Parse the metadata content returned by the Ollama server.
    
    Parameters:
        content (str): Metadata content in the specified XML-like format.
        verbose (bool): Enable verbose logging.
    
    Returns:
        dict or None: Dictionary containing author, year, title, and language or None if parsing fails.
    """
    title_match = re.search(r'<TITLE>(.*?)</TITLE>', content, re.DOTALL)
    year_match = re.search(r'<YEAR>(\d{4})</YEAR>', content, re.DOTALL)
    author_match = re.search(r'<AUTHOR>(.*?)</AUTHOR>', content, re.DOTALL)
    language_match = re.search(r'<LANGUAGE>(.*?)</LANGUAGE>', content, re.DOTALL)
    logging.debug(f"Parsing metadata: {content}")

    # If TITLE tag is incomplete, try to match it differently
    if not title_match:
        title_match = re.search(r'TITLE>(.*?)</TITLE>', content, re.DOTALL)
    
    # Include additional checks for None values
    title = title_match.group(1).strip() if title_match else None
    author = author_match.group(1).strip() if author_match else None
    year = year_match.group(1).strip() if year_match else "Unknown"
    language = language_match.group(1).strip().lower() if language_match else "en"
    
    if not title_match:
        logging.warning(f"No match for title in {content}.")
        return None
    if not author_match:
        logging.warning(f"No match for author in {content}.")
        return None
    if not year_match:
        logging.info(f"No match for year in {content}. Assigning 'Unknown'.")
        year = "Unknown"
    else:
        year = sanitize_filename(year_match.group(1).strip())
    if not language_match:
        logging.info(f"No match for language in {content}. Attempting to detect language from text.")
        language_match = detect_language(content)
        language = language_match.group(1).strip().lower() if language_match else "en"
        logging.info(f"Assuming language is {language}.")
    else:
        language = sanitize_filename(language_match.group(1).strip()).lower()
    
    title = sanitize_filename(title_match.group(1).strip())
    author = author_match.group(1).strip()
    
    # Sanitize filenames
    title = sanitize_filename(title) if title else "unknown"
    author = sanitize_filename(author) if author else "unknown"
    year = sanitize_filename(year)
    language = sanitize_filename(language)
    
    if any(placeholder in (title.lower(), author.lower(), year.lower(), language.lower()) for placeholder in ["unknown", "UnknownAuthor", "n a", ""]):
        logging.warning("Warning: Found 'unknown', 'n a', or empty strings in title, year, author, or language.")
        logging.warning(f"Parsed metadata - Title: '{title}', Year: '{year}', Author: '{author}', Language: '{language}'")
        return None
    return {'author': author, 'year': year, 'title': title, 'language': language}

# Communication with Ollama server
def send_to_ollama_server(text, filename, openai_client, max_attempts=4, verbose=False):
    attempt = 1
    while attempt <= max_attempts and not shutdown_flag.is_set():
        logging.debug(f"Consulting Ollama server on file: {filename} (Attempt: {attempt})")
        base_filename = os.path.splitext(os.path.basename(filename))[0]
        prompt = (
            f"Extract the first author name (ignore other authors), year of publication, title, and language from the following text, considering the filename '{base_filename}' which may contain clues. "
            f"I need the output **only** in the following format with no additional text or explanations: \n"
            f"<TITLE>The publication title</TITLE> \n<YEAR>2023</YEAR> \n<AUTHOR>Lastname Surname</AUTHOR> \n<LANGUAGE>en</LANGUAGE> \n\n"
            f"Here is the extracted text:\n{text}"
        )
        messages = [{"role": "user", "content": prompt}]
        
        if verbose:
            logging.debug(f"Sending prompt to Ollama server: {prompt}")
        
        try:
            response = openai_client.chat.completions.create(
                model=MODEL_NAME,
                temperature=0.7,
                max_tokens=250,
                messages=messages,
                timeout=300  # Set a timeout of 300 seconds
            )
        except Exception as e:
            logging.error(f"Error communicating with Ollama server for {filename}: {e}")
            if attempt < max_attempts:
                wait_time = 2 ** attempt
                logging.info(f"Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                attempt += 1
                continue
            return ""
        if verbose:
            logging.debug(f"Raw content received from server: (start){response.choices[0].message.content}(end)")
        output = response.choices[0].message.content.strip()
        if verbose:
            logging.debug(f"Metadata content received from server: {output}")
        if ("lastname" in output.lower() or "surname" in output.lower()) and attempt < max_attempts:
            logging.warning(f"{filename} - Detected placeholders in response. Retrying...")
            attempt += 1
            time.sleep(2 ** attempt)
            continue
        # Updated Validation: Use re.search to ensure all tags are present anywhere in the response
        title = re.search(r'<TITLE>(.*?)</TITLE>', output, re.DOTALL)
        year = re.search(r'<YEAR>(\d{4})</YEAR>', output, re.DOTALL)
        author = re.search(r'<AUTHOR>(.*?)</AUTHOR>', output, re.DOTALL)
        language = re.search(r'<LANGUAGE>(.*?)</LANGUAGE>', output, re.DOTALL)
        
        if title and year and author and language:
            if verbose:
                logging.debug(f"The output contains all four fields: title, year, author, and language.")
            return output
        else:
            logging.warning(f"Unexpected response format from Ollama server: {output}.")
            return output
    logging.error(f"{filename} - Maximum retry attempts reached for sending to Ollama server.")
    return ""


def send_metadata_to_ollama(metadata_info, filename, openai_client, max_attempts=4, verbose=False):
    attempt = 1
    while attempt <= max_attempts and not shutdown_flag.is_set():
        logging.debug(f"Sending file metadata to Ollama server for file: {filename} (Attempt: {attempt})")
        metadata_str = "\n".join([f"{key}: {value}" for key, value in metadata_info.items()])
        prompt = (
            f"Unable to extract meaningful text from the file '{os.path.splitext(os.path.basename(filename))[0]}'. "
            f"Here is the file metadata that might contain clues:\n{metadata_str}\n\n"
            f"Based on this metadata, extract the first author name (ignore other authors), year of publication, title of the publication, and language. "
            f"Provide the output in the following format:\n"
            f"<TITLE>The publication title</TITLE>\n"
            f"<YEAR>2023</YEAR>\n"
            f"<AUTHOR>Lastname Surname</AUTHOR>\n"
            f"<LANGUAGE>en</LANGUAGE>"
        )
        messages = [{"role": "user", "content": prompt}]
        logging.debug(f"Sending prompt to Ollama server: {prompt}")
        
        try:
            response = openai_client.chat.completions.create(
                model=MODEL_NAME,
                temperature=0.7,
                max_tokens=250,
                messages=messages
            )
        except Exception as e:
            logging.error(f"Error sending metadata to Ollama server for {filename}: {e}")
            if attempt < max_attempts:
                wait_time = 2 ** attempt
                logging.info(f"Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                attempt += 1
                continue
            return ""
        output = response.choices[0].message.content.strip()
        logging.debug(f"Raw content received from server: (start){output}(end)")
        if ("lastname" in output.lower() or "surname" in output.lower()) and attempt < max_attempts:
            logging.warning("Detected placeholders in response. Retrying...")
            attempt += 1
            time.sleep(2 ** attempt)
            continue
        # Extract all required fields
        title = re.search(r'<TITLE>(.*?)</TITLE>', output, re.DOTALL)
        year = re.search(r'<YEAR>(\d{4})</YEAR>', output, re.DOTALL)
        author = re.search(r'<AUTHOR>(.*?)</AUTHOR>', output, re.DOTALL)
        language = re.search(r'<LANGUAGE>(.*?)</LANGUAGE>', output, re.DOTALL)
        
        if title and year and author and language:
            if verbose:
                logging.debug(f"The output contains all four fields: title, year, author, and language.")
            return output
        else:
            logging.warning(f"Unexpected response format from Ollama server: {output}.")
            return output
    logging.error("Maximum retry attempts reached for sending metadata to Ollama server.")
    return ""


# Filename sanitization to ensure safe filesystem operations
def sanitize_filename(name):
    name = re.sub(r'[\\/*?:"<>|]', "", name)
    if " " in name:
        name = re.sub(r'(\w)\.', r'\1', name)
        parts = name.split()
        new_parts = []
        current_initials = ""
        for part in parts:
            if len(part) == 1:
                current_initials += part
            else:
                if current_initials:
                    new_parts.append(current_initials)
                    current_initials = ""
                new_parts.append(part)
        if current_initials:
            new_parts.append(current_initials)
        name = " ".join(new_parts)
    return name.strip().replace('/', '')


# Execute the generated rename commands script
def execute_rename_commands(script_path):
    try:
        # Ensure the file is closed before executing
        with open(script_path, 'r') as script_file:
            pass  # Just to ensure it's accessible and can be opened
        subprocess.run(['bash', script_path], check=True)
        logging.info(f"Successfully executed rename commands from {script_path}")
    except subprocess.CalledProcessError as e:
        logging.error(f"Error executing rename commands: {e}. Please check the rename script for issues.")
    except FileNotFoundError:
        logging.error(f"Rename script '{script_path}' not found.")
    except PermissionError:
        logging.error(f"Permission denied while executing the rename script '{script_path}'. Ensure it is executable.")
    except Exception as e:
        logging.error(f"Unexpected error during rename command execution: {e}")


# Main processing function for each file
def process_file(file_path, openai_client, max_attempts=MAX_RETRY_ATTEMPTS, verbose=False, tools_available=None, optional_libraries=None):
    """
    Process a single file to extract metadata and prepare rename commands.
    
    Parameters:
        file_path (str): Path to the file to process.
        openai_client: Initialized OpenAI client for communication.
        max_attempts (int): Maximum number of processing attempts.
        verbose (bool): Enable verbose logging.
        tools_available (dict): Dictionary indicating availability of tools.
        optional_libraries (dict): Dictionary indicating availability of optional libraries.
    
    Returns:
        dict or None: Extracted metadata dictionary or None if processing fails.
    """
    # Create a mutable copy of the read-only tools_available
    if tools_available is not None and isinstance(tools_available, MappingProxyType):
        tools_available = dict(tools_available)
    
    # Check global failure counters and disable tools if needed
    if 'kraken_failure_count' in globals() and globals()['kraken_failure_count'] >= 3:
        if tools_available and tools_available.get('kraken', False):
            logging.warning("Kraken has failed too many times. Disabling for this session.")
            tools_available['kraken'] = False
    
    # Same for other tools if needed
    if 'tesseract_failure_count' in globals() and globals()['tesseract_failure_count'] >= 3:
        if tools_available and tools_available.get('tesseract', False):
            logging.warning("Tesseract has failed too many times. Disabling for this session.")
            tools_available['tesseract'] = False
    
    if tools_available is None:
        tools_available = {}
    if optional_libraries is None:
        optional_libraries = {}

    logging.debug(f"Processing file: {file_path}")
    if shutdown_flag.is_set():
        logging.info(f"Shutdown initiated. Skipping processing for {file_path}.")
        return None
    
    # Check if file is empty
    if is_file_empty(file_path):
        logging.error(f"File {file_path} is empty. Skipping processing.")
        return None
    
    if file_path.lower().endswith('.pdf') and (tools_available.get('PyPDF2') or tools_available.get('pdfminer')):
        if not is_pdf_valid(file_path, tools_available=tools_available, optional_libraries=optional_libraries):
            logging.warning(f"Invalid PDF file detected: {file_path}. Skipping text extraction.")
            with file_lock:
                with open("unparseables.lst", "a") as unparseable_file:
                    unparseable_file.write(f"{file_path}\n")
                    unparseable_file.flush()
            return None  # Skip further processing for this file
    
    # Initialize OCR attempts tracking
    ocr_attempts = set()

    attempt = 0
    while attempt < max_attempts and not shutdown_flag.is_set():
        attempt += 1
        try:
            logging.info(f"Starting processing for: {file_path} (Attempt {attempt})")
            
            text = ""
            # Initialize extraction methods based on file extension
            extraction_methods = []
            if file_path.lower().endswith('.pdf'):
                extraction_methods = ['pdf_extractor', 'ocr']
            elif file_path.lower().endswith('.epub'):
                extraction_methods = ['ebooklib', 'textract']
            elif file_path.lower().endswith('.docx'):
                extraction_methods = ['docx', 'textract']
            elif file_path.lower().endswith(('.azw', '.azw3')):
                extraction_methods = ['ebooklib', 'mobi', 'textract']
            elif file_path.lower().endswith('.mobi'):
                extraction_methods = ['mobi', 'textract']
            elif file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.tiff', '.tif')):
                extraction_methods = ['ocr']
            elif file_path.lower().endswith(('.doc', '.xls', '.xlsx', '.ppt', '.pptx', '.odt', '.ods', '.odp', '.html', '.xml', '.rtf', '.md', '.txt')):
                extraction_methods = ['textract']
            else:
                logging.warning(f"Unsupported file format for {file_path}. Skipping.")
                return None

            extraction_attempts = 0
            for method in extraction_methods:
                if shutdown_flag.is_set():
                    logging.info(f"Shutdown initiated. Stopping extraction for {file_path}.")
                    return None
                extraction_attempts += 1
                logging.debug(f"Attempting text extraction using {method} for {file_path}")
                if method == 'fitz':
                    text = extract_text_with_fitz(file_path, verbose=verbose)
                elif method == 'pdfplumber':
                    text = extract_text_with_pdfplumber(file_path, verbose=verbose)
                elif method == 'pdfminer':
                    text = extract_text_with_pdfminer(file_path, verbose=verbose)
                elif method == 'pdf_extractor':
                        text = extract_text_from_pdf(
                        file_path,
                        perform_ocr_if_needed=False,
                        ocr_available={'tesseract': tools_available.get('tesseract', False), 'kraken': tools_available.get('kraken', False)},
                        verbose=verbose,
                        tools_available=tools_available,
                        ocr_attempts=ocr_attempts,
                    )
                        
                elif method == 'ocr':
                    logging.info(f"Attempting OCR on {file_path}")
                    if file_path.lower().endswith('.pdf'):
                        # Try Tesseract first, then fall back to Kraken if needed
                        text = perform_ocr_on_first_pages(
                            file_path,
                            num_pages=3,
                            verbose=verbose,
                            tools_available=tools_available,
                            ocr_attempts=ocr_attempts,
                        )
                    else:
                        image = None
                        if tools_available.get('pillow', False):
                            try:
                                image = Image.open(file_path)
                            except Exception as e:
                                logging.error(f"Error opening image {file_path} for OCR: {e}")
                        
                        if image:
                            # Try Tesseract first, regardless of OCR_PREFERRED_ORDER
                            if tools_available.get('tesseract', False) or shutil.which('tesseract'):
                                try:
                                    # Try using pytesseract if available
                                    if optional_libraries.get('pytesseract', False):
                                        import pytesseract
                                        lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
                                        if not lang:
                                            lang = 'eng'  # Default to English if no languages found
                                        logging.debug(f"Performing OCR on image with pytesseract using languages: {lang}")
                                        text = pytesseract.image_to_string(image, lang=lang)
                                        if text.strip() and is_text_meaningful(text, file_path=file_path):
                                            logging.debug(f"Successfully extracted meaningful text using pytesseract ({len(text)} characters)")
                                            if verbose:
                                                logging.debug(f"Extracted text: {text[:500]}")
                                    
                                    # If pytesseract failed or isn't available, use subprocess
                                    if not text:
                                        # Save image to temp file and use tesseract command
                                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                                            temp_path = temp_file.name
                                        
                                        image.save(temp_path)
                                        logging.debug(f"Saved image to temporary file {temp_path}")
                                        
                                        lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
                                        if not lang:
                                            lang = 'eng'  # Default to English if no languages found
                                        
                                        cmd = ['tesseract', temp_path, 'stdout', '-l', lang]
                                        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                                        text = result.stdout.strip()
                                        
                                        # Clean up the temporary file
                                        try:
                                            os.unlink(temp_path)
                                        except:
                                            pass
                                        
                                        if text.strip() and is_text_meaningful(text, file_path=file_path):
                                            logging.debug(f"Successfully extracted meaningful text using tesseract subprocess ({len(text)} characters)")
                                            if verbose:
                                                logging.debug(f"Extracted text: {text[:500]}")
                                except Exception as e:
                                    logging.error(f"Error performing OCR with Tesseract on {file_path}: {e}")
                            
                            # Only try Kraken if Tesseract failed
                            if not text:
                                text = perform_ocr_on_image(
                                    image,
                                    ocr_available={'tesseract': tools_available.get('tesseract', False), 'kraken': tools_available.get('kraken', False)},
                                    tools_available=tools_available,
                                    optional_libraries=optional_libraries,
                                    verbose=verbose,
                                    ocr_attempts=ocr_attempts,
                                )
                        else:
                            # If we couldn't load the image with Pillow, try direct Tesseract subprocess
                            if shutil.which('tesseract'):
                                try:
                                    lang = '+'.join([lang for lang in REQUIRED_TESS_LANGS if lang in get_installed_tesseract_langs()])
                                    if not lang:
                                        lang = 'eng'  # Default to English if no languages found
                                    
                                    cmd = ['tesseract', file_path, 'stdout', '-l', lang]
                                    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                                    text = result.stdout.strip()
                                    
                                    if text.strip() and is_text_meaningful(text, file_path=file_path):
                                        logging.debug(f"Successfully extracted meaningful text using tesseract subprocess directly on file ({len(text)} characters)")
                                        if verbose:
                                            logging.debug(f"Extracted text: {text[:500]}")
                                except Exception as e:
                                    logging.error(f"Error performing OCR with Tesseract directly on {file_path}: {e}")
                elif method == 'ebooklib':
                    if file_path.lower().endswith('.epub'):
                        text = extract_text_from_epub(
                            file_path,
                            textract_available=tools_available.get('textract', False),
                            tools_available=tools_available,
                            verbose=verbose
                        )
                    elif file_path.lower().endswith(('.azw', '.azw3')):
                        text = extract_text_from_azw(
                            file_path,
                            verbose=verbose,
                            tools_available=tools_available
                        )
                elif method == 'docx':
                    text = extract_text_from_docx(
                        file_path,
                        verbose=verbose,
                        tools_available=tools_available
                    )
                elif method == 'mobi':
                    text = extract_text_from_mobi(
                        file_path,
                        verbose=verbose,
                        tools_available=tools_available
                    )
                elif method == 'textract':
                    text = extract_text_with_textract(file_path, verbose=verbose)
                else:
                    logging.warning(f"Extraction method {method} not implemented for {file_path}")
                    continue

                if text and is_text_meaningful(text, file_path=file_path):
                    logging.debug(f"Text extraction successful using {method}")
                    break  # Successful extraction
                else:
                    logging.debug(f"Text extraction using {method} did not yield meaningful text.")
                    text = ""  # Reset text for next method

            if not text:
                logging.warning(f"All extraction methods failed to yield meaningful text for {file_path}. Attempting to extract metadata.")
                metadata_info = extract_file_metadata(file_path, verbose=verbose, tools_available=tools_available, optional_libraries=optional_libraries)
                if not metadata_info:
                    logging.warning(f"No metadata extracted for {file_path}. Cannot proceed.")
                    with file_lock:
                        with open("unparseables.lst", "a") as unparseable_file:
                            unparseable_file.write(f"{file_path}\n")
                            unparseable_file.flush()
                    return None
                metadata_content = send_metadata_to_ollama(metadata_info, file_path, openai_client, verbose=verbose)
            else:
                metadata_content = send_to_ollama_server(text, file_path, openai_client, verbose=verbose)

            if metadata_content:
                metadata = parse_metadata(metadata_content, verbose=verbose)
                if metadata:
                    # Correct author names and detect language
                    corrected_author = sort_author_names(metadata['author'], openai_client, verbose=verbose)
                    metadata['author'] = corrected_author
                    if not valid_author_name(metadata['author']):
                        logging.warning(f"Author name '{metadata['author']}' is not valid.")
                    return metadata
                else:
                    logging.warning(f"Metadata parsing failed for {file_path}. Retrying...")
            else:
                logging.warning(f"No metadata content received for {file_path}. Retrying...")
        except Exception as e:
            logging.error(f"Error processing {file_path}: {e}. Retrying...")
        logging.info(f"Retrying processing for {file_path} (Attempt {attempt}/{max_attempts})...")
    logging.error(f"Failed to process {file_path} after {max_attempts} attempts.")
    with file_lock:
        try:
            with open("unparseables.lst", "a") as unparseable_file:
                unparseable_file.write(f"{file_path}\n")
                unparseable_file.flush()
        except Exception as e:
            logging.error(f"Failed to write to unparseables.lst: {e}")
    return None


# Escape special characters in filenames for shell commands
def escape_special_chars(filename):
    """
    Safely escape special characters in filenames for shell commands.
    
    Tries to use shlex.quote for comprehensive escaping.
    Falls back to regex-based escaping if shlex is unavailable.
    
    Parameters:
        filename (str): The filename to escape.
    
    Returns:
        str: Escaped filename.
    """
    try:
        import shlex
        return shlex.quote(filename)
    except ImportError:
        logging.warning("shlex module not available. Falling back to regex-based escaping.")
        return re.sub(r'([$`"\\])', r'\\\1', filename)

def detect_language(text, file_path, tools_available=None, optional_libraries=None):
    """
    Detect the language of the given text using langdetect.
    If langdetect is not available or fails, infer language from metadata or use heuristics.
    """
    global LANGDETECT_AVAILABLE

    if tools_available is None:
        tools_available = {}
    if optional_libraries is None:
        optional_libraries = {}

    if LANGDETECT_AVAILABLE:
        try:
            language = detect(text)
            logging.debug(f"Detected language using langdetect: {language}")
            return language
        except Exception as e:
            logging.error(f"Language detection with langdetect failed for {file_path}: {e}")

    # Fallback to metadata-based detection
    metadata = extract_file_metadata(file_path, verbose=False, tools_available=tools_available, optional_libraries=optional_libraries)
    language = metadata.get('language')
    if language:
        logging.debug(f"Detected language from metadata: {language}")
        return language.lower()

    # Heuristic detection based on common words
    logging.debug("Attempting heuristic language detection.")
    common_words = {
        'en': ['the', 'and', 'is', 'to', 'are', 'in', 'it', 'you', 'that', 'he'],
        'de': ['der', 'die', 'und', 'zu', 'ist', 'nicht', 'ich', 'es', 'du', 'das'],
        'fr': ['le', 'les', 'des', 'et', 'la', 'est', 'pas', 'je', 'il', 'vous'],
        'es': ['el', 'la', 'y', 'en', 'que', 'los', 'para', 'es', 'una', 'por'],
        'ar': ['الذي', 'التي', 'في', 'من', 'على', 'مع', 'كان', 'هذا', 'هل', 'نعم'],
        'lat': ['est', 'sunt', 'erat', 'qui', 'quae', 'ex', 'ad', 'in', 'et', 'non'],
        # Add more languages and common words as needed
    }
    text_lower = text.lower()
    for lang_code, words in common_words.items():
        if any(word in text_lower for word in words):
            logging.debug(f"Heuristically detected language as '{lang_code}' using common words.")
            return lang_code

    # Fallback default
    logging.debug(f"Defaulting language to 'en' for {file_path}")
    return 'en'


def extract_file_metadata(file_path, verbose=False, tools_available=None, optional_libraries=None):
    if tools_available is None:
        tools_available = {}
    if optional_libraries is None:
        optional_libraries = {}

    logging.debug(f"Extracting file metadata for: {file_path}")
    metadata = {}
    try:
        if optional_libraries.get('magic', False):
            file_type = magic.from_file(file_path, mime=True)
            metadata['file_type'] = file_type
        else:
            metadata['file_type'] = "Unknown"

        metadata['file_size'] = os.path.getsize(file_path)

        creation_time = os.path.getctime(file_path)
        modification_time = os.path.getmtime(file_path)
        metadata['creation_time'] = datetime.fromtimestamp(creation_time).isoformat()
        metadata['modification_time'] = datetime.fromtimestamp(modification_time).isoformat()

        # Extract PDF metadata if the file is a PDF
        if file_path.lower().endswith('.pdf'):
            pdf_info_extracted = False
            if optional_libraries.get('PyPDF2', False):
                try:
                    with open(file_path, 'rb') as f:
                        reader = PdfReader(f)
                        pdf_info = reader.metadata
                        if pdf_info:
                            metadata['title'] = pdf_info.title
                            metadata['author'] = pdf_info.author
                            metadata['subject'] = pdf_info.subject
                            metadata['creator'] = pdf_info.creator
                            metadata['producer'] = pdf_info.producer
                            metadata['creation_date'] = pdf_info.creation_date
                            metadata['modification_date'] = pdf_info.modification_date
                            pdf_info_extracted = True
                    logging.debug(f"Extracted PDF metadata using PyPDF2: {metadata}")
                except Exception as e:
                    logging.error(f"Error extracting PDF metadata with PyPDF2 from {file_path}: {e}")
            if not pdf_info_extracted and optional_libraries.get('pdfminer', False):
                try:
                    # pdfminer can extract metadata using PdfDocument
                    from pdfminer.pdfdocument import PDFDocument
                    from pdfminer.pdfparser import PDFParser
                    with open(file_path, 'rb') as f:
                        parser = PDFParser(f)
                        doc = PDFDocument(parser)
                        if doc.info:
                            metadata.update({k.decode('utf-8'): v.decode('utf-8') for k, v in doc.info[0].items()})
                            pdf_info_extracted = True
                    logging.debug(f"Extracted PDF metadata using pdfminer: {metadata}")
                except Exception as e:
                    logging.error(f"Error extracting PDF metadata with pdfminer from {file_path}: {e}")
            if not pdf_info_extracted and tools_available.get('textract', False):
                try:
                    # Textract may not extract metadata, but we can try to extract some info
                    text = textract.process(file_path).decode('utf-8')
                    metadata['extracted_text'] = text[:2000]
                    logging.debug(f"Extracted text using textract for metadata extraction.")
                except Exception as e:
                    logging.error(f"Error extracting text from {file_path} using textract for metadata: {e}")
            if not pdf_info_extracted:
                logging.warning(f"Could not extract PDF metadata for {file_path}")
        logging.debug(f"Extracted metadata: {metadata}")

    except Exception as e:
        logging.error(f"Error extracting metadata from {file_path}: {e}")

    return metadata

def prompt_install_optional_libraries(optional_libraries, no_install=False):
    """
    Prompt the user to install missing optional libraries.

    Parameters:
        optional_libraries (dict): Dictionary indicating availability of optional libraries.
        no_install (bool): If True, do not attempt to install even if prompted.

    Returns:
        None
    """
    missing_libs = [lib for lib, available in optional_libraries.items() if not available and lib in ['fitz', 'pdfplumber']]
    for lib in missing_libs:
        install_commands = []
        if lib == 'fitz':
            install_commands = ['pip', 'install', 'PyMuPDF']
        elif lib == 'pdfplumber':
            install_commands = ['pip', 'install', 'pdfplumber']

        if install_commands:
            if prompt_install_tool(lib, install_commands, auto_install=True, no_install=no_install):
                # Attempt to import again after installation
                try:
                    if lib == 'fitz':
                        import fitz
                    elif lib == 'pdfplumber':
                        import pdfplumber
                    optional_libraries[lib] = True
                    logging.info(f"{lib} installed and available.")
                except ImportError:
                    optional_libraries[lib] = False
                    logging.error(f"Failed to import {lib} even after installation.")

def disable_tool_on_failures(tool_name, failure_counter, max_failures=3):
    """
    Disable a tool after it fails too many times.
    
    Parameters:
        tool_name (str): The name of the tool in tools_available dict
        failure_counter (int): Current failure count
        max_failures (int): Maximum allowed failures before disabling
        
    Returns:
        bool: True if tool should be disabled, False otherwise
    """
    if failure_counter >= max_failures:
        logging.warning(f"{tool_name} has failed {failure_counter} times. Disabling for this session.")
        return True
    return False

# Main function to process all files in the directory
def main(directory, verbose, force, singletask, debug, NO_INSTALL):
    global optional_libraries
    global LANGDETECT_AVAILABLE
    rename_script_path = "rename_commands.sh"

    progress_bar = None  # Initialize progress_bar to None
    
    # Initialize external_tools dictionary with all relevant tools
    external_tools = {
        'tesseract': False,
        'kraken': False, 
        'djvutxt': False,
        'textract': essential_libraries.get('textract', False)
    }

    # Check for Tesseract availability before Kraken to ensure fallback works
    external_tools['tesseract'] = check_tesseract_availability()

    # Attempt to import langdetect
    try:
        from langdetect import detect, DetectorFactory
        DetectorFactory.seed = 0  # For consistent results across runs
        LANGDETECT_AVAILABLE = True
        logging.debug("langdetect is available for language detection.")
    except ImportError:
        LANGDETECT_AVAILABLE = False
        logging.warning("langdetect is not installed. Language detection will rely on file names and metadata.")

    # Check for Kraken availability
    # Initialize global counter for Kraken failures
    kraken_failure_count = 0

    kraken_available = False
    if shutil.which('kraken'):
        # Try quick command version check first
        try:
            result = run_subprocess_with_timeout(['kraken', '--version'], timeout=10, capture_output=True, text=True)
            if result and result.returncode == 0:
                logging.debug(f"Kraken command-line tool is available: {result.stdout.strip()}")
                
                # The command exists and runs - that's enough to try the command-line approach
                kraken_available = True
                external_tools['kraken'] = True
            else:
                logging.warning("Kraken command-line tool check failed")
        except Exception as e:
            logging.warning(f"Error checking kraken command-line: {e}")
        
        # If command-line check failed, or as additional validation, check Python modules
        if not kraken_available:
            try:
                # Check if we can import the necessary modules
                import kraken
                import kraken.binarization
                import kraken.pageseg
                
                # Try multiple possible module paths for recognition capabilities
                recognition_available = False
                try:
                    # Try newer API first (kraken 3.x+)
                    from kraken.recognition import load_any
                    recognition_available = True
                    logging.debug("Kraken 3.x recognition module is available.")
                except ImportError:
                    try:
                        # Try rpred API (kraken 2.x)
                        import kraken.rpred
                        recognition_available = True
                        logging.debug("Kraken 2.x rpred module is available.")
                    except ImportError:
                        logging.warning("Neither Kraken recognition nor rpred modules could be imported.")
                
                if recognition_available:
                    kraken_available = True
                    external_tools['kraken'] = True
                    logging.debug("Kraken OCR is available and all required modules are importable.")
                else:
                    logging.warning("Kraken OCR core modules are available but recognition modules are missing.")
                    logging.warning("Will attempt to repair Kraken installation if needed.")
            except ImportError as e:
                logging.warning(f"Kraken OCR is installed but Python modules are not properly importable: {e}")
                logging.warning("Will attempt to repair Kraken installation if needed.")
                kraken_available = False
    else:
        logging.warning("Kraken OCR is not installed or not found in PATH. Kraken OCR functionality will be skipped.")
        kraken_available = False

    if not kraken_available and not NO_INSTALL:
        # Prompt user to install Kraken with PDF support
        install_commands = ['pip', 'install', '--force-reinstall', 'kraken[pdf]']  # Force reinstall with PDF extras
        if prompt_install_tool('Kraken OCR', install_commands, auto_install=True, no_install=NO_INSTALL):
            if shutil.which('kraken'):
                try:
                    # Verify installation
                    import kraken
                    import kraken.binarization
                    import kraken.pageseg
                    
                    # Check for recognition modules
                    recognition_available = False
                    try:
                        from kraken.recognition import load_any
                        recognition_available = True
                    except ImportError:
                        try:
                            import kraken.rpred
                            recognition_available = True
                        except ImportError:
                            logging.warning("Kraken recognition modules not found after installation.")
                    
                    if recognition_available:
                        external_tools['kraken'] = True
                        logging.info("Kraken OCR installed successfully with all required modules.")
                        
                        # After successful installation, download required models
                        required_models = [            # Example model IDs, adjust based on your needs
                            '10.5281/zenodo.2577813',  # en_best.mlmodel
                            '10.5281/zenodo.11113737', # reichenau_lat_cat_099218.mlmodel
                            '10.5281/zenodo.7631619',  # cremma-generic-1.0.1.mlmodel
                            '10.5281/zenodo.7050296',  # arabic_best.mlmodel
                            '10.5281/zenodo.10519596', # german_print.mlmodel
                            '10.5281/zenodo.7933402',  # austriannewspapers.mlmodel
                            # Add other model IDs as needed
                        ]
                        for model_id in required_models:
                            try:
                                logging.info(f"Downloading Kraken model: {model_id}")
                                model_result = run_subprocess_with_timeout(['kraken', 'get', model_id], timeout=180, check=False)
                                if model_result and model_result.returncode == 0:
                                    logging.info(f"Successfully downloaded model: {model_id}")
                                else:
                                    logging.warning(f"Failed to download Kraken model {model_id}")
                            except Exception as e:
                                logging.error(f"Unexpected error while downloading Kraken model {model_id}: {e}")
                    else:
                        logging.error("Kraken OCR installation was incomplete. Missing recognition modules.")
                        external_tools['kraken'] = False
                except ImportError as e:
                    logging.error(f"Kraken OCR installation was incomplete. Missing modules: {e}")
                    external_tools['kraken'] = False
            else:
                logging.error("Kraken OCR installation failed or 'kraken' is not in PATH.")
                external_tools['kraken'] = False

    
    # Check for pyvips (required for Kraken PDF support)
    pyvips_available = False
    try:
        import pyvips
        pyvips_available = True
        logging.debug("pyvips is available.")
    except Exception as e:
        logging.warning(f"pyvips import failed: {e}")
        logging.warning("Kraken PDF support requires pyvips to function properly.")
        
        # Determine if libvips is installed
        libvips_installed = False
        try:
            result = subprocess.run(['pkg-config', '--exists', 'vips'], check=False)
            libvips_installed = result.returncode == 0
        except:
            pass
        
        if not libvips_installed and shutil.which('brew') and not NO_INSTALL:
            try:
                user_input = input("libvips is not found but is required for Kraken. Install with Homebrew? [y/N]: ").strip().lower()
                if user_input == 'y':
                    logging.info("Installing libvips with Homebrew...")
                    subprocess.run(['brew', 'install', 'vips'], check=True)
                    logging.info("Successfully installed libvips. Setting up environment variables...")
                    
                    # Get libvips location
                    vips_lib_path = subprocess.run(['brew', '--prefix', 'vips'], capture_output=True, text=True).stdout.strip() + "/lib"
                    os.environ['DYLD_LIBRARY_PATH'] = vips_lib_path + ":" + os.environ.get('DYLD_LIBRARY_PATH', '')
                    
                    # Try installing pyvips if it's not already
                    logging.info("Installing/reinstalling pyvips...")
                    subprocess.run(['pip', 'uninstall', '-y', 'pyvips'], check=False)
                    subprocess.run(['pip', 'install', '--no-cache-dir', 'pyvips'], check=True)
                    
                    # We can't reliably import pyvips now because the dynamic loader cache is already populated
                    # Instead, give instructions to the user
                    logging.info("libvips and pyvips installed. Please restart the script in a new terminal session.")
                    logging.info(f"Before running, add this to your shell profile (~/.zshrc or ~/.bash_profile):")
                    logging.info(f"    export DYLD_LIBRARY_PATH=\"{vips_lib_path}:$DYLD_LIBRARY_PATH\"")
                    
                    if input("Would you like to exit now to restart with the new environment? [y/N]: ").strip().lower() == 'y':
                        logging.info("Exiting for restart with new environment variables.")
                        sys.exit(0)
            except Exception as e:
                logging.error(f"Failed to install libvips or pyvips: {e}")
        
        logging.info("Continuing without Kraken PDF support...")

    # Check for Tesseract availability
    if shutil.which('tesseract'):
        external_tools['tesseract'] = True
        logging.debug("Tesseract OCR is available.")
    else:
        logging.warning("Tesseract OCR is not installed or not found in PATH. Tesseract OCR functionality will be skipped.")
        # Prompt user to install Tesseract
        current_os = platform.system()
        install_commands = []
        if current_os == "Darwin":
            install_commands = ['brew', 'install', 'tesseract']
        elif current_os == "Linux":
            if shutil.which('apt-get'):
                install_commands = ['sudo', 'apt-get', 'install', '-y', 'tesseract-ocr']
            elif shutil.which('yum'):
                install_commands = ['sudo', 'yum', 'install', '-y', 'tesseract']
        
        if install_commands:
            if prompt_install_tool('Tesseract OCR', install_commands, auto_install=True, no_install=NO_INSTALL):
                external_tools['tesseract'] = True
                if not shutil.which('tesseract'):
                    logging.error("Tesseract OCR installation failed or 'tesseract' is not in PATH.")
                    external_tools['tesseract'] = False
        else:
            logging.error("Unsupported OS or package manager. Please install Tesseract OCR manually.")

    # Check for djvutxt availability
    external_tools['djvutxt'] = shutil.which('djvutxt') is not None
    if not external_tools['djvutxt']:
        logging.warning("djvutxt is not installed or not found in PATH. DJVU files will be skipped.")
    else:
        logging.debug("djvutxt is available.")

    if not os.path.exists(directory):
        logging.error("The specified directory does not exist.")
        sys.exit(1)
    
    files = [f for f in os.listdir(directory) if f.lower().endswith(SUPPORTED_EXTENSIONS)]
    if not files:
        logging.info("No supported files found in the specified directory.")
        sys.exit(0)
    
    try:
        with file_lock:
            with open(rename_script_path, "w") as mv_file:
                mv_file.write("#!/bin/bash\n")
                mv_file.write('set -e\n')
                mv_file.flush()
    
        # Dynamically find tessdata directories and set TESSDATA_PREFIX
        tessdata_dirs = find_tessdata_dirs()

        if tessdata_dirs:
            TESSDATA_PREFIX = tessdata_dirs[0]  # Use the first found tessdata directory
            os.environ['TESSDATA_PREFIX'] = TESSDATA_PREFIX
            logging.debug(f"Set TESSDATA_PREFIX to: {TESSDATA_PREFIX}")
        else:
            logging.warning("Could not find 'tessdata' directory. Please set the TESSDATA_PREFIX environment variable manually.")
            TESSDATA_PREFIX = None  # Handle cases where tessdata is not found

        # Additions for 'fitz' and 'pdfplumber'
        tools_available = MappingProxyType({
            'tesseract': external_tools.get('tesseract', False),
            'kraken': external_tools.get('kraken', False),
            'djvutxt': external_tools.get('djvutxt', False),
            'textract': essential_libraries.get('textract', False),
            'pdfminer': optional_libraries.get('pdfminer', False),
            'PyPDF2': optional_libraries.get('PyPDF2', False),
            'pillow': optional_libraries.get('pillow', False),
            'ebooklib': optional_libraries.get('ebooklib', False),
            'beautifulsoup4': optional_libraries.get('beautifulsoup4', False),
            'docx': optional_libraries.get('docx', False),
            'mobi': optional_libraries.get('mobi', False),
            'pytesseract': optional_libraries.get('pytesseract', False),
            'pdf2image': optional_libraries.get('pdf2image', False),
            'fitz': optional_libraries.get('fitz', False),
            'pdfplumber': optional_libraries.get('pdfplumber', False)
        })

        # Log tool availability
        tool_messages = {
            'tesseract': "Tesseract OCR is available.",
            'kraken': "Kraken OCR is available.",
            'djvutxt': "djvutxt is available.",
            'textract': "textract is available.",
            'pdfminer': "pdfminer is available.",
            'PyPDF2': "PyPDF2 is available.",
            'pillow': "Pillow (PIL) is available.",
            'ebooklib': "ebooklib is available.",
            'beautifulsoup4': "beautifulsoup4 is available.",
            'docx': "python-docx is available.",
            'mobi': "mobi library is available.",
            'pytesseract': "pytesseract is available.",
            'pdf2image': "pdf2image is available.",
            'fitz': "PyMuPDF (fitz) is available.",
            'pdfplumber': "pdfplumber is available.",
        }

        for tool, available in tools_available.items():
            if available:
                logging.debug(tool_messages.get(tool, f"{tool} is available."))
            else:
                logging.warning(tool_messages.get(tool, f"{tool} is not available."))

        
        prompt_install_optional_libraries(optional_libraries, no_install=NO_INSTALL)
        
        # Check and install missing Tesseract languages
        if tools_available['tesseract']:
            check_and_install_tesseract_langs(REQUIRED_TESS_LANGS)

        openai_client = get_openai_client()

        if singletask:
            max_workers = 1
            logging.info("Single-task mode enabled. Processing files sequentially.")
        else:
            max_workers = min(32, os.cpu_count() + 4)
            logging.info(f"Processing files with {max_workers} concurrent workers.")
        
        # Initialize ProgressBar
        
        progress_bar = ProgressBar(total=len(files), desc="Processing Files", enable=optional_libraries.get('tqdm', False))
        # Initialize ProgressBar
        
        if singletask:
            # Sequential processing
            for filename in files:
                if shutdown_flag.is_set():
                    logging.info("Shutdown flag detected. Stopping further processing.")
                    break
                file_path = os.path.join(directory, filename)
                try:
                    metadata = process_file(file_path, openai_client, max_attempts=MAX_RETRY_ATTEMPTS, verbose=verbose, tools_available=tools_available, optional_libraries=optional_libraries)
                    if metadata:
                        author = metadata['author']
                        title = metadata['title']
                        year = metadata['year']
                        
                        if not author or not title:
                            logging.warning(f"Missing author or title for {file_path}. Skipping renaming.")
                            with file_lock:
                                with open("unparseables.lst", "a") as unparseable_file:
                                    unparseable_file.write(f"{file_path}\n")
                                    unparseable_file.flush()
                            continue
                        first_author = sanitize_filename(author.split(", ")[0])
                        target_dir = os.path.join(directory, first_author)
                        if year == "Unknown":
                            logging.info(f"Year is unknown for {file_path}. Assigning 'UnknownYear' in filename.")
                            year = "UnknownYear"
                        else:
                            year = sanitize_filename(year)
                        sanitized_title = sanitize_filename(title)
                        file_extension = filename.split('.')[-1]
                        new_filename = f"{year} {sanitized_title}.{file_extension}"
                        new_file_path = os.path.join(target_dir, new_filename)
                        escaped_file_path = escape_special_chars(file_path)
                        escaped_target_dir = escape_special_chars(target_dir)
                        escaped_new_file_path = escape_special_chars(new_file_path)
                        with file_lock:
                            with open(rename_script_path, "a") as mv_file:
                                mv_file.write(f"mkdir -p {escaped_target_dir}\n")
                                mv_file.write(f"mv {escaped_file_path} {escaped_new_file_path}\n")
                                mv_file.flush()
                        logging.debug(f"Rename command added for: {file_path}")
                    else:
                        with file_lock:
                            with open("unparseables.lst", "a") as unparseable_file:
                                unparseable_file.write(f"{file_path}\n")
                                unparseable_file.flush()
                        logging.info(f"Added to unparseables: {file_path}")
                except Exception as e:
                    logging.error(f"Failed to process {file_path}: {e}")
                    with file_lock:
                        with open("unparseables.lst", "a") as unparseable_file:
                            unparseable_file.write(f"{file_path}\n")
                            unparseable_file.flush()
                finally:
                    progress_bar.update(1)
        else:
            # Concurrent processing
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {}
                
                # Submit all tasks
                for filename in files:
                    future = executor.submit(
                        process_with_timeout,  # Wrap the process_file with timeout
                        process_file, 
                        os.path.join(directory, filename), 
                        openai_client,
                        max_attempts=MAX_RETRY_ATTEMPTS, 
                        verbose=verbose, 
                        tools_available=tools_available,
                        optional_libraries=optional_libraries
                    )
                    future_to_file[future] = filename
                
                # Create a timeout for the whole processing batch
                overall_deadline = time.time() + (300 * len(files))  # Average 5 minutes per file
                
                # Process completed tasks
                try:
                    for future in concurrent.futures.as_completed(future_to_file, timeout=max(300, len(files) * 10)):
                        # Check if we need to shutdown
                        if shutdown_flag.is_set() or time.time() > overall_deadline:
                            logging.info("Shutdown flag detected or deadline reached. Breaking out of the processing loop.")
                            for fut in list(future_to_file.keys()):
                                if not fut.done():
                                    fut.cancel()
                            break
                        
                        # Process the completed future
                        filename = future_to_file[future]
                        file_path = os.path.join(directory, filename)
                        try:
                            metadata = future.result()
                            # (rest of your existing processing code)
                            
                        except Exception as e:
                            logging.error(f"Failed to process {file_path}: {e}")
                            with file_lock:
                                with open("unparseables.lst", "a") as unparseable_file:
                                    unparseable_file.write(f"{file_path}\n")
                                    unparseable_file.flush()
                        finally:
                            progress_bar.update(1)
                except concurrent.futures.TimeoutError:
                    logging.error("Overall batch processing timed out")
                    for fut in list(future_to_file.keys()):
                        if not fut.done():
                            fut.cancel()
    
        progress_bar.close()

        with file_lock:
            try:
                os.chmod(rename_script_path, 0o755)
                logging.debug(f"Set execute permissions for {rename_script_path}.")
            except Exception as e:
                logging.error(f"Error setting execute permissions for {rename_script_path}: {e}")
        logging.info(f"\nRename commands have been written to {rename_script_path}")
        logging.info("Please review the commands before executing them.")
        
        if force:
            logging.info("Force option detected. Executing rename commands...")
            execute_rename_commands(rename_script_path)
        else:
            logging.info(f"To execute the rename commands, run: bash {rename_script_path}")

    except KeyboardInterrupt:
        logging.info("KeyboardInterrupt received. Initiating graceful shutdown...")
        shutdown_flag.set()
        if progress_bar and progress_bar.enable:
            progress_bar.close()
    except Exception as e:
        logging.critical(f"Unhandled exception in main workflow: {e}", exc_info=True)
        sys.exit(1)
    finally:
        if progress_bar and progress_bar.enable:
            progress_bar.close()
        logging.info("Script terminated.")

        # Stop the logging queue listener
        logger = logging.getLogger()
        if hasattr(logger, 'queue_listener'):
            logger.queue_listener.stop()

# ------------------------------------------------
# Definitions and imports
# (placed here, so that logging is possible)

# Import pdfminer if available
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    from pdfminer.pdfparser import PDFSyntaxError
    optional_libraries_pdfminer = True
    logging.debug("pdfminer is available for PDF extraction.")
except ImportError:
    optional_libraries_pdfminer = False
    logging.warning("pdfminer.six is not installed. PDFMiner-based extraction will be skipped.")
except Exception as e:
    optional_libraries_pdfminer = False
    logging.error(f"Unexpected error importing pdfminer: {e}")
    logging.warning("PDFMiner-based extraction will be skipped.")

# Import optional libraries with warnings if not available
optional_libraries = {
    'tqdm': TQDM_AVAILABLE,
    'ebooklib': False,
    'beautifulsoup4': False,
    'docx': False,
    'mobi': False,
    'magic': False,
    'PyPDF2': False,
    'pdf2image': False,
    'pytesseract': False,
    'pdfminer': optional_libraries_pdfminer,
    'pillow': False 
}

# Attempt to import OpenAI client
try:
    from openai import OpenAI
except ImportError:
    logging.critical("OpenAI client library is not installed. Please install it using 'pip install openai'.")
    sys.exit(1)

# Define a shutdown flag for graceful termination
shutdown_flag = threading.Event()

# Signal handler to set the shutdown flag
def signal_handler(signum, frame):
    """
    Signal handler for SIGINT, SIGTERM, and other interrupt signals.
    Sets the shutdown flag and logs the interrupt event.
    """
    if not shutdown_flag.is_set():  # Only log once
        signal_name = {
            signal.SIGINT: "SIGINT (Ctrl+C)",
            signal.SIGTERM: "SIGTERM",
            signal.SIGTSTP: "SIGTSTP (Ctrl+Z)"
        }.get(signum, f"Signal {signum}")
        
        logging.info(f"Received {signal_name}. Initiating graceful shutdown...")
        shutdown_flag.set()
        
        # When SIGINT (Ctrl+C) is received, make sure to set a proper exit code
        if signum == signal.SIGINT:
            # Instead of exit, use a delayed exit to allow threads to clean up
            def exit_after_timeout():
                time.sleep(MAX_WAIT_AFTER_SHUTDOWN)
                if threading.main_thread().is_alive():
                    logging.warning("Forced exit after timeout - some threads may not have terminated cleanly")
                    os._exit(1)  # Force exit if threads aren't responding
            
            # Start the force-exit timer
            threading.Thread(target=exit_after_timeout, daemon=True).start()
    
    # If SIGTSTP (Ctrl+Z), convert it to a graceful shutdown instead of suspending
    if signum == signal.SIGTSTP:
        # Prevent the default behavior (job suspension)
        signal.signal(signal.SIGTSTP, signal.SIG_IGN)

# Register signal handlers for SIGINT and SIGTERM
signal.signal(signal.SIGINT, signal_handler)
signal.signal(signal.SIGTERM, signal_handler)
signal.signal(signal.SIGTSTP, signal_handler)  # Handle Ctrl+Z

# Attempt to import essential and optional libraries and set availability flags
essential_libraries = {
    'textract': False,
    'openai': True  # Already imported above
}

# Import textract
try:
    import textract
    essential_libraries['textract'] = True
    logging.debug("textract is available for text extraction.")
except ImportError:
    logging.warning("textract is not installed. PDF extraction will fallback to PyPDF2 if available.")

# Attempt to import Pillow (PIL)
try:
    from PIL import Image
    optional_libraries['pillow'] = True
    logging.debug("Pillow (PIL) is available for image processing.")
except ImportError:
    optional_libraries['pillow'] = False
    logging.warning("Pillow (PIL) is not installed. Image processing will be skipped or limited.")

try:
    from ebooklib import epub
    from ebooklib import ITEM_DOCUMENT as EPUB_ITEM_DOCUMENT
    optional_libraries['ebooklib'] = True
    logging.debug("ebooklib is available for EPUB processing.")
except ImportError:
    logging.warning("ebooklib is not installed. EPUB and MOBI file processing will be skipped.")

try:
    from bs4 import BeautifulSoup
    optional_libraries['beautifulsoup4'] = True
    logging.debug("beautifulsoup4 is available for HTML parsing.")
except ImportError:
    logging.warning("beautifulsoup4 is not installed. HTML and EPUB text extraction will be limited.")

try:
    import docx
    optional_libraries['docx'] = True
    logging.debug("python-docx is available for DOCX processing.")
except ImportError:
    logging.warning("python-docx is not installed. DOCX file processing will be skipped.")

try:
    import mobi
    optional_libraries['mobi'] = True
    logging.debug("mobi library is available for MOBI processing.")
except ImportError:
    logging.warning("mobi library is not installed. MOBI and AZW file processing will be skipped.")

try:
    import magic
    optional_libraries['magic'] = True
    logging.debug("python-magic is available for MIME type detection.")
except ImportError:
    logging.warning("python-magic is not installed. MIME type detection will be skipped.")

try:
    from PyPDF2 import PdfReader
    from PyPDF2.errors import PdfReadError
    optional_libraries['PyPDF2'] = True
    logging.debug("PyPDF2 is available for PDF extraction.")
except ImportError:
    logging.warning("PyPDF2 is not installed. PDF files will be skipped.")
    optional_libraries['PyPDF2'] = False

try:
    from pdf2image import convert_from_path
    optional_libraries['pdf2image'] = True
    logging.debug("pdf2image is available for converting PDFs to images.")
except ImportError:
    logging.warning("pdf2image is not installed. OCR on specific pages will be skipped.")
    optional_libraries['pdf2image'] = False

try:
    import pytesseract
    optional_libraries['pytesseract'] = True
    logging.debug("pytesseract is available for OCR.")
except ImportError:
    logging.warning("pytesseract is not installed. OCR functionality will be limited.")
    optional_libraries['pytesseract'] = False

try:
    import fitz  # PyMuPDF
    optional_libraries['fitz'] = True
    logging.debug("PyMuPDF (fitz) is available for PDF extraction.")
except ImportError:
    optional_libraries['fitz'] = False
    logging.warning("PyMuPDF (fitz) is not installed. PDF extraction using PyMuPDF will be skipped.")

try:
    import pdfplumber
    optional_libraries['pdfplumber'] = True
    logging.debug("pdfplumber is available for PDF extraction.")
except ImportError:
    optional_libraries['pdfplumber'] = False
    logging.warning("pdfplumber is not installed. PDF extraction using pdfplumber will be skipped.")



# Entry Point
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Parse various document formats to extract metadata and generate rename commands."
    )
    parser.add_argument(
        "directory", 
        nargs='?', 
        default=".", 
        help="Directory containing document files (default: current directory)"
    )
    parser.add_argument(
        "--verbose", 
        action="store_true", 
        help="Enable verbose output (detailed logs in logfile)"
    )

    parser.add_argument(
        "--no-install", 
        action="store_true", 
        help="Do not prompt for installing missing tools. Proceed without them."
    )

    parser.add_argument(
        "--debug", 
        action="store_true", 
        help="Enable debug mode: detailed logs on stdout only, no logfile"
    )
    parser.add_argument(
        "--force", 
        action="store_true", 
        help="Automatically execute rename commands"
    )
    parser.add_argument(
        "--singletask", 
        action="store_true", 
        help="Disable concurrent processing and process files sequentially"
    )
    args = parser.parse_args()
    DIRECTORY = args.directory
    VERBOSE = args.verbose
    DEBUG = args.debug
    FORCE = args.force
    SINGLETASK = args.singletask
    NO_INSTALL = args.no_install
    setup_logging(VERBOSE, DEBUG)

    main(DIRECTORY, VERBOSE, FORCE, SINGLETASK, DEBUG, NO_INSTALL)
